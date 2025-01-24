# -*- coding: utf-8 -*-
"""
Created on Wed Nov 20 18:17:29 2024

@author: MottaRS
"""


#imports necessários
# import os
import docx
# import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt
from PIL import Image
import re
from datetime import date

# from docx2pdf import convert

#parâmetros globais
fonte_name = 'Times New Roman'
tamanho_fonte_titulo = 15
tamanho_fonte_texto = 12
tamanho_fonte_tabela = 10
tamanho_fonte_legenda = 10
estilo_tabela = 'Medium Grid 1 Accent 1'
estilo_tabela = 'Medium Shading 1 Accent 1'
estilo_tabela = 'Light Grid Accent 1'


def configuraTexto(doc, texto,  titulo = False, iten=0):

    if titulo:
        # Seção
        # doc.section ?
        texto_formatado=texto
        if iten:
            num_topico = doc.sections_num
            while len(num_topico)<iten:
                num_topico.append(0)
            num_topico = num_topico[:iten]
            num_topico[-1]=num_topico[-1]+1
            doc.sections_num = num_topico

            # Cria a string formatada
            numero_topico = '.'.join(map(str, num_topico))
            texto_formatado = f"{numero_topico}.  {texto}"
            print( texto_formatado)
            
        paragrafo = doc.add_paragraph(texto_formatado)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragrafo.runs[0].font.name = fonte_name
        paragrafo.runs[0].font.size = Pt(tamanho_fonte_titulo)
        paragrafo.runs[0].font.bold = True
    else:
        # NORMAL
        paragrafo = doc.add_paragraph(texto)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragrafo.runs[0].font.name = fonte_name
        paragrafo.runs[0].font.size = Pt(tamanho_fonte_texto)
        
    return doc

def identificacaoDoRelatorio(doc,  relatorio = None, projeto = None, objetivo = None, Sistema = 'PACAT'):

    if relatorio is None:
        relatorio = u'Avaliação Integridade de Duto Baseado em Relatório de Inspeção .'
        # relatorio = ' '
    if projeto is None:
        projeto = u'Avaliação de Integridade Estrutural de Dutos Terrestres.'
        # projeto = ' '
    if objetivo is None:
        objetivo = u'Este relatório tem como principal objetivo apresentar os resultados das análises dos defeitos, obtidos com o sistema {Sistema}.'
        # objetivo = ' '

    titulo = u'Relatório de Análise de Integridade do Duto & Manual de Manutenção'

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #RELATÓRIO
    paragrafo2 = doc.add_paragraph(u'Duto: xxxx')
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo2.runs[0].font.bold = True
    paragrafo2.add_run(relatorio)
    for run in paragrafo2.runs:
        run.font.name = fonte_name
        run.font.size = Pt(tamanho_fonte_texto)

    #PROJETO
    paragrafo3 = doc.add_paragraph('Trecho: xxxx')
    paragrafo3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo3.runs[0].font.bold = True
    paragrafo3.add_run(projeto)
    for run in paragrafo3.runs:
        run.font.name = fonte_name
        run.font.size = Pt(tamanho_fonte_texto)

    #OBJETIVO
    paragrafo4 = doc.add_paragraph('Engenheiro Responsável: xxxx\n')
    paragrafo4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo4.runs[0].font.bold = True
    paragrafo4.add_run(objetivo)
    for run in paragrafo4.runs:
        run.font.name = fonte_name
        run.font.size = Pt(tamanho_fonte_texto)

    return doc

def get_section_level(text):
    """
    Verifica se a string é uma seção, retorna o nível da seção e o texto sem os números.
    
    Args:
        text (str): Texto a ser avaliado.
    
    Returns:
        tuple: (nivel_da_secao, texto_sem_numeracao)
               Retorna (0, texto_original) se não for uma seção.
    """
    
    match = re.match(r'^(\d+(\.\d+)*)(.*)', text.strip())
    if match:
        level = match.group(1).count('.') + 1
        cleaned_text = match.group(3).strip()
        return level, cleaned_text
    return 0, text.strip()
    

def Criar_Relatorio(file_name = 'Relatorio_de_Inspecao.docx', insp=[]):
    ##############################################################################
    ######## INICIO ########
    ##############################################################################
    
    #organizadores gerais do relatório
    num_tabela = 1
    num_figura = 1
    
    #criação do documento
    doc = docx.Document()
    #creating sections_num list handeling by function "configuraTexto"
    doc.sections_num=[0]
    
    #inicialização dos dados de identificação do relatório
    #identificacaoDoRelatorio(doc,  relatorio = None, projeto = None, objetivo = None)
    doc = identificacaoDoRelatorio(doc)
    
    doc = configuraTexto(doc, 'Informações Gerais', True, 1)
    njoint = len(insp.df_joints)
    
    texto = [
    'Duto: [Inserir nome ou código do duto]',
    'Trecho: [Inserir o trecho analisado]',
    f'Data da análise: {date.today()}',
    f'Inspeção atual: {insp.date}',
    'Inspeção anterior: [Informar ano e método de ILI, se aplicável]',
    f'Comprimento total analisado: {insp.df_joints.S_pos.iloc[-1]} m',
    f'Número de juntas analisadas: {njoint}',
    f'Total de anomalias detectadas: {len(insp.df_Def)}'
    ]
    
    for txt in texto:
        doc = configuraTexto(doc, txt)
    
    
    texto = ['Resumo da metodologia usada: ',
    'Explicação breve sobre a análise realizada no pipe tally.',
    'Abordagem usada (determinística e/ou probabilística).',
    'Ferramentas e modelos aplicados no PACAT.']
    
    doc = configuraTexto(doc, 'Introdução', True, 1)
    for txt in texto:
        doc = configuraTexto(doc, txt)
    
    doc = configuraTexto(doc, 'Descrição do duto', True, 2)
    doc = configuraTexto(doc, '[incluir texto sobre o duto]')
        
    fign=0
    fig_list = [
    'anomalies_histogram.png',
    'Depth_histogram.png',
    'Defects_Sizes.png',
    'Geo_Loc.png',
    'Joint_Position.png',
    'Defects_Clock_Position.png',
    'Defects_Clusters.png',
    'Defects_ERF.png',
    'Defects_ASSESSMENT.png',
    'Defects_Critical_Depth.png',
    'Defects_Critical_Size.png',
    'Defects_MSOP.png',
    'Current Inspection.png',
    'Features_position.png',
    # 'Future Assessment (5 years).png',
    #'Metal_Loss_position.png'
    ]
    
    
    legenda = 'asdsdgn.yhou'
    print('certo legenda_2')
    paragrafo1 = doc.add_paragraph(legenda)
    paragrafo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo1.runs[0].font.name = 'Times New Roman'
    print('certo legenda_3')
    paragrafo1.runs[0].font.size = Pt(tamanho_fonte_legenda)
    
    # for img in fig_list:
    #     paragrafo = doc.add_paragraph()
    #     run_imagem = paragrafo.add_run()
    #     run_imagem.add_picture(img, width=Cm(16))
    
    textos=[
        "2 Análise Geral de Anomalias",
        "2.1 Distribuição Geral",
        "Este gráfico apresenta a distribuição das categorias ou comentários atribuídos às anomalias detectadas na inspeção. Ele facilita a identificação das principais causas e características das falhas no pipe tally, permitindo uma análise inicial das condições do duto.",
        "Figura 1: Histograma das categorias de anomalias detectadas (Figura: anomalies_histogram.png).",
        "2.2 Profundidade e Dimensão dos Defeitos",
        "Este histograma mostra a frequência de profundidades das anomalias ao longo do duto. Essa informação é essencial para identificar a severidade das perdas metálicas e definir prioridades de reparo.",
        "Figura 2: Histograma das profundidades detectadas (Figura: Depth_histogram.png).",
        "Este gráfico apresenta os tamanhos longitudinais dos defeitos ao longo do duto. Essa análise ajuda a identificar regiões de maior deterioração e avaliar os impactos potenciais na integridade estrutural.",
        "Figura 3: Distribuição dos tamanhos dos defeitos ao longo do duto (Figura: Defects_Sizes.png).",
        "3 Localização e Caracterização dos Defeitos",
        "3.1 Localização Geográfica",
        "Este mapa mostra a posição geográfica das anomalias detectadas. Ele é fundamental para associar os defeitos às condições ambientais e operacionais específicas de cada trecho do duto.",
        "Figura 4: Localização georreferenciada dos defeitos (Figura: Geo_Loc.png).",
        "3.2 Posição ao Longo do Duto",
        "Este gráfico ilustra a localização dos defeitos ao longo dos tubos em relação às juntas. Ele ajuda a avaliar se os defeitos estão concentrados em áreas críticas, como as interfaces entre tubos.",
        "Figura 5: Posição dos defeitos em relação às juntas (Figura: Joint_Position.png).",
        "Este gráfico detalha a posição angular e a dimensão dos defeitos. A visualização ajuda a identificar padrões de corrosão ou falhas associadas a fatores específicos, como fluxo ou orientação do duto.",
        "Figura 6: Posição circunferencial e dimensão dos defeitos ao longo do duto (Figura: Defects_Clock_Position.png).",
        "3.3 Identificação de Clusters",
        "Este gráfico identifica clusters de defeitos interagentes ao longo do duto. Essas interações podem aumentar a criticidade dos defeitos e requerem atenção especial durante a avaliação de integridade.",
        "Figura 7: Identificação dos defeitos interagentes (Figura: Defects_Clusters.png).",
        "4 Avaliação Atual de Integridade",
        "4.1 Análise dos Fatores de Reparo (ERF)",
        "Este gráfico mostra os fatores de reparo calculados para os defeitos detectados. O ERF é uma métrica que indica a severidade de cada defeito em relação aos limites operacionais do duto.",
        "Figura 8: Fatores de reparo (ERF) dos defeitos ao longo do duto (Figura: Defects_ERF.png).",
        "Esta figura apresenta uma análise visual detalhada dos ERFs ao longo do duto, ajudando a identificar as regiões mais críticas para priorizar ações de manutenção.",
        "Figura 9: Análise detalhada dos fatores de reparo por posição no duto (Figura: Defects_ASSESSMENT.png).",
        "4.2 Limites Críticos",
        "Este gráfico compara as profundidades detectadas com os valores críticos que resultariam em ERF = 1. Ele é usado para determinar quais defeitos já atingiram ou estão próximos do limite crítico.",
        "Figura 10: Profundidade crítica dos defeitos para ERF = 1 (Figura: Defects_Critical_Depth.png).",
        "Similar à análise de profundidade crítica, este gráfico foca no comprimento dos defeitos em relação aos valores que resultariam em ERF = 1, destacando os mais severos.",
        "Figura 11: Comprimento crítico dos defeitos para ERF = 1 (Figura: Defects_Critical_Size.png).",
        "4.3 Pressão Máxima de Operação Segura",
        "Este gráfico apresenta a pressão máxima de operação segura (MSOP) calculada para cada defeito. Ele ajuda a determinar se o duto pode continuar operando em condições seguras.",
        "Figura 12: Pressão máxima segura para cada defeito (Figura: Defects_MSOP.png).",
        "4.4 Probabilidade de Falha",
        "Esta figura mostra a probabilidade de falha (PF) associada a cada defeito, juntamente com os fatores de reparo (ERFs). Ela oferece uma visão abrangente da criticidade de cada anomalia.",
        "Figura 13: Probabilidade de falha e ERF combinados (Figura: Current Inspection.png).",
        #"5 Avaliação Futura",
        #"5.1 Projeção de Crescimento",
        #"Este gráfico apresenta a projeção da integridade do duto com base nas taxas de crescimento de corrosão. Ele permite prever a evolução dos defeitos e planejar intervenções preventivas.",
        #"Figura 14: Avaliação futura considerando um horizonte de 5 anos (Figura: Future Assessment (5 years).png).",
        "6 Conclusão",
        "Resumo das condições do duto.",
        "Identificação de pontos críticos e recomendações para reparo ou monitoramento.",
        "7 Apêndice",
        "Metodologias aplicadas em detalhe.",
        "Tabelas suplementares, se necessário."
    ]
    
    for txt in textos:
        iten, tt=get_section_level(txt)
        if iten:
            doc = configuraTexto(doc, tt,True, iten)
        
        else:
            if txt.startswith('Fig'):
                paragrafo = doc.add_paragraph()
                run_imagem = paragrafo.add_run()
                run_imagem.add_picture(fig_list[fign], width=Cm(16))
                fign+=1
            doc = configuraTexto(doc, tt)
        
    
    doc.save(file_name)