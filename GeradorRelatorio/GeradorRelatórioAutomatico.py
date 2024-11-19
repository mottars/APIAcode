#imports necess�rios
import os
import docx
import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt
from PIL import Image

#par�metros globais
tamanho_fonte_titulo = 15
tamanho_fonte_texto = 12
tamanho_fonte_tabela = 10
tamanho_fonte_legenda = 10
estilo_tabela = 'Medium Grid 1 Accent 1'
estilo_tabela = 'Medium Shading 1 Accent 1'
estilo_tabela = 'Light Grid Accent 1'

#fun��es auxiliares
def configuraTabela(doc, tabela, linhas, colunas, titulo_unico,  altura_linha = None, legenda = None):
	if altura_linha is None:
		altura_linha = 0.6
	for i in range(0, linhas):
		tabela.rows[i].height = Cm(altura_linha)
		for j in range(0, colunas):
			if tabela.rows[i].cells[j].text != '':
				tabela.rows[i].cells[j].width = Cm(7)
				tabela.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
				try:
					tabela.rows[i].cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
				except:
					print('erro')
				try:
					tabela.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(tamanho_fonte_tabela)
				except:
					print('erro 2')
				tabela.rows[i].cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	if titulo_unico:
		cell00 = tabela.columns[0].cells[0]
		cell01 = tabela.columns[1].cells[0]
		cell00.merge(cell01)
		
	if legenda is not None:
		paragrafo1 = doc.add_paragraph(legenda)
		paragrafo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
		paragrafo1.runs[0].font.name = 'Times New Roman'
		paragrafo1.runs[0].font.size = Pt(tamanho_fonte_legenda)

	return tabela	

def configuraImagem(doc, arquivo,  legenda = None):

	paragrafo = doc.add_paragraph()
	paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run_imagem = paragrafo.add_run()
	run_imagem.add_picture(arquivo)
	print('certo legenda')
	if legenda is not None:
		print('certo legenda_2')
		paragrafo1 = doc.add_paragraph(legenda)
		paragrafo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
		paragrafo1.runs[0].font.name = 'Times New Roman'
		print('certo legenda_3')
		paragrafo1.runs[0].font.size = Pt(tamanho_fonte_legenda)
	print('certo saiu')
	return doc

def configuraTexto(doc, texto,  titulo = False):

	if titulo:
		paragrafo = doc.add_paragraph(texto)
		paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
		paragrafo.runs[0].font.name = 'Times New Roman'
		paragrafo.runs[0].font.size = Pt(tamanho_fonte_titulo)
		paragrafo.runs[0].font.bold = True
	else:
		paragrafo = doc.add_paragraph(texto)
		paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		paragrafo.runs[0].font.name = 'Times New Roman'
		paragrafo.runs[0].font.size = Pt(tamanho_fonte_texto)

	return doc

#fun��es construtoras do documento
#quantidadeDeCasos e identificacaoDoRelatorio s�o chamadas apenas uma vez!
#o restante das fun��es s�o chamadas uma vez para cada caso que for compor o relat�rio!
def quantidadeDeCasos(doc, num_topico, num_tabela,  qtd_casos = None):


	titulo = str(num_topico) + u'. QUANTIDADE DE CASOS'
	texto = u'	Neste relat�rio, ser� analisada a quantidade de casos de defeitos de corros�o que consta especificada na Tabela 1 a seguir.'
	texto1 = u'QUANTIDADE DE CASOS'
	legenda = u'Tabela ' + str(num_tabela) + u': Quantidade de casos analisados'
	num_tabela += 1

	if qtd_casos is None:
		qtd_casos = 1

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO
	doc = configuraTexto(doc, texto)

	paragrafo1 = doc.add_paragraph(legenda)
	paragrafo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragrafo1.runs[0].font.name = 'Times New Roman'
	paragrafo1.runs[0].font.size = Pt(tamanho_fonte_legenda)

	#TABELA
	tabela = doc.add_table(rows = 1, cols = 2)
	tabela.style = estilo_tabela
	tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
	linha_tabela = tabela.rows[0].cells #primeira linha da tabela como uma lista
	tabela.rows[0].height = Cm(0.6)
	linha_tabela[0].text = texto1 #elemento (0, 0) da tabela
	linha_tabela[1].text = str(qtd_casos) #elemento (0, 1) da tabela
	for item in linha_tabela:
		item.width = Cm(5)
		item.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
		item.paragraphs[0].runs[0].font.name = 'Times New Roman'
		item.paragraphs[0].runs[0].font.size = Pt(tamanho_fonte_tabela)
		item.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	

	return doc, num_tabela

def identificacaoDoRelatorio(doc, num_topico,  relatorio = None, projeto = None, objetivo = None):

	if relatorio is None:
		relatorio = u'An�lise de Dutos com Defeitos Isolados.'
		relatorio = ' '
	if projeto is None:
		projeto = u'Avalia��o de Integridade Estrutural de Dutos Terrestres.'
		projeto = ' '
	if objetivo is None:
		objetivo = u'Este relat�rio tem como principal objetivo apresentar os resultados das an�lises de dutos com defeitos isolados, obtidos com o sistema CORDUT.'
		objetivo = ' '

	titulo = str(num_topico) + u'. IDENTIFICA��O DO RELAT�RIO'

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#RELAT�RIO
	paragrafo2 = doc.add_paragraph(u'Relat�rio: ')
	paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	paragrafo2.runs[0].font.bold = True
	paragrafo2.add_run(relatorio)
	for run in paragrafo2.runs:
		run.font.name = 'Times New Roman'
		run.font.size = Pt(tamanho_fonte_texto)

	#PROJETO
	paragrafo3 = doc.add_paragraph('Projeto: ')
	paragrafo3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	paragrafo3.runs[0].font.bold = True
	paragrafo3.add_run(projeto)
	for run in paragrafo3.runs:
		run.font.name = 'Times New Roman'
		run.font.size = Pt(tamanho_fonte_texto)

	#OBJETIVO
	paragrafo4 = doc.add_paragraph('Objetivo: \n')
	paragrafo4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	paragrafo4.runs[0].font.bold = True
	paragrafo4.add_run(objetivo)
	for run in paragrafo4.runs:
		run.font.name = 'Times New Roman'
		run.font.size = Pt(tamanho_fonte_texto)

	return doc

def modeloDoDuto(doc, num_topico, num_subtopico, num_tabela,  diametro_externo = None, espessura = None, comprimento = None, taxa_ovalizacao = None, nome_modelo = None):

	titulo = str(num_topico) + '. CASO ' + str(num_topico - 2)

	titulo1 = str(num_topico) + '.' + str(num_subtopico) + ' MODELO DO DUTO'
	texto = [u'	As principais caracter�sticas geom�tricas dos segmentos de duto (di�metro Externo ', 'D', 'e', ', comprimento ', 'L', 'duto',
                 ', espessura de parede ', 't',
                 u') analisados est�o descritas na Tabela 2. Al�m disso, caracter�sticas referente � ovaliza��o inicial da se��o transversal do segmento de duto onde se encontr�o defeito de corros�o e ao nome do modelo analisado tamb�m est�o contidas nesta tabela.']

	legenda = 'Tabela ' + str(num_tabela) + u': Caracter�sticas geom�tricas do duto.'
	num_tabela += 1

	#T�TULOS
	doc = configuraTexto(doc, titulo, titulo = True)
	doc = configuraTexto(doc, titulo1, titulo = True)

	#TEXTO
	paragrafo2 = doc.add_paragraph()
	paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	for i in range(len(texto)):
		run = paragrafo2.add_run(texto[i])
		if i in [1, 2, 4, 5, 7]:
			run.font.italic = True
			if i == 2 or i== 5:
				run.font.subscript = True
		run.font.name = 'Times New Roman'
		run.font.size = Pt(tamanho_fonte_texto)

	#TABELA
	if diametro_externo is None:
		diametro_externo = 600
	if espessura is None:
		espessura = 10
	if comprimento is None:
		comprimento = 3000
	if taxa_ovalizacao is None:
		taxa_ovalizacao = 5
	if nome_modelo is None:
		nome_modelo = 'DT1'

	paragrafo2 = doc.add_paragraph(legenda)
	paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragrafo2.runs[0].font.name = 'Times New Roman'
	paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)
	
	tabela = doc.add_table(rows = 6, cols = 2)
	tabela.style = estilo_tabela
	tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
	tabela.rows[0].cells[0].text = u'DADOS DO SEGMENTO DO DUTO'
##	tabela.rows[1].cells[0].text = u'Di�metro Externo: De (mm)'
	lista_de_tab = [u'Di�metro Externo (','D', 'e', ') (mm)']
	paragrafo_celula = tabela.rows[1].cells[0].add_paragraph()
	paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
	for i in range(len(lista_de_tab)):
		run = paragrafo_celula.add_run(lista_de_tab[i])
		if i in [1, 2]:
			run.font.italic = True
			if i == 2:
				run.font.subscript = True
		run.font.name = 'Times New Roman'
##	tabela.rows[2].cells[0].text = u'Espessura do Duto t (mm)'
	lista_de_tab = [u'Espessura do Duto (','t', ') (mm)']
	paragrafo_celula = tabela.rows[2].cells[0].add_paragraph()
	paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
	for i in range(len(lista_de_tab)):
		run = paragrafo_celula.add_run(lista_de_tab[i])
		if i in [1]:
			run.font.italic = True
##			if i == 2:
##				run.font.subscript = True
		run.font.name = 'Times New Roman'

##	tabela.rows[3].cells[0].text = u'Comprimento do Duto Lduto (mm)'
	lista_de_tab = [u'Comprimento do Duto (','L', 'duto', ') (mm)']
	paragrafo_celula = tabela.rows[3].cells[0].add_paragraph()
	paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
	for i in range(len(lista_de_tab)):
		run = paragrafo_celula.add_run(lista_de_tab[i])
		if i in [1, 2]:
			run.font.italic = True
			if i == 2:
				run.font.subscript = True
		run.font.name = 'Times New Roman'
##	tabela.rows[4].cells[0].text = u'Taxa de Ovaliza��o fo (%)'
	lista_de_tab = [u'Taxa de Ovaliza��o (','f', 'o', ') (%)']
	paragrafo_celula = tabela.rows[4].cells[0].add_paragraph()
	paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
	for i in range(len(lista_de_tab)):
		run = paragrafo_celula.add_run(lista_de_tab[i])
		if i in [1, 2]:
			run.font.italic = True
			if i == 2:
				run.font.subscript = True
		run.font.name = 'Times New Roman'
	tabela.rows[5].cells[0].text = u'Nome do Modelo'
	#tabela.rows[0].cells[1].text = ''
	tabela.rows[1].cells[1].text = str(diametro_externo)#str(format(diametro_externo, ".2e"))
	tabela.rows[2].cells[1].text = str(espessura)#str(format(espessura, ".2e"))
	tabela.rows[3].cells[1].text = str(comprimento)#str(format(comprimento, ".2e"))
	tabela.rows[4].cells[1].text = str(taxa_ovalizacao)#str(format(taxa_ovalizacao, ".2f"))
	tabela.rows[5].cells[1].text = nome_modelo

	tabela = configuraTabela(doc, tabela, 6, 2, titulo_unico = True)#, legenda = legenda)

	return doc, num_tabela

def propriedadesDoMaterial(doc, num_topico, num_subtopico, num_tabela, num_figura,  arquivo_grafico = None, modulo_elasticidade = None, coeficiente_poisson = None, tensao_escoamento = None, tensao_ruptura = None, tensoes = None, deformacoes = None):

	titulo = str(num_topico) + '.' + str(num_subtopico) + ' PROPRIEDADES DO MATERIAL'
	texto = [u'	As propriedades do material do duto encontram-se na Tabela 3. A partir desses dados, foi constru�da a curva tens�o verdadeira ', 'versus',
                 u' deforma��o total verdadeira do material utilizada nas an�lises, indicada na Figura 1 e os seus pontos est�o apresentados na Tabela 4. A mesma resulta da combina��o de um trecho linear da origem (0,0) at� a tens�o de escoamento, mais um trecho n�o linear definido segundo a equa��o de Ramberg-Osgood.']
	if arquivo_grafico is None:
		arquivo_grafico = 'tensao_versus_deformacao.png'

	legenda1 = 'Figura ' + str(num_figura) + u': Curva tens�o deforma��o verdadeira do material.'
	num_figura += 1
	legenda2 = 'Tabela ' + str(num_tabela) + ': Propriedades do material do duto.'
	legenda3 = 'Tabela ' + str(num_tabela + 1) + u': Entrada de dados (tens�o/deforma��o) da curva verdadeira do material.'
	num_tabela += 2

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO
	paragrafo2 = doc.add_paragraph()
	paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	for i in range(len(texto)):
		run = paragrafo2.add_run(texto[i])
		if i == 1:
			run.font.italic = True
		run.font.name = 'Times New Roman'
		run.font.size = Pt(tamanho_fonte_texto)

	paragrafo2 = doc.add_paragraph(legenda2)
	paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragrafo2.runs[0].font.name = 'Times New Roman'
	paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)
	
	#TABELA 1 - PROPRIEDADES DO MATERIAL
	if modulo_elasticidade is None:
		modulo_elasticidade = 200000
	if coeficiente_poisson is None:
		coeficiente_poisson = 0.3
	if tensao_escoamento is None:
		tensao_escoamento = 534.1
	if tensao_ruptura is None:
		tensao_ruptura = 754.18
	tabela1 = doc.add_table(rows = 5, cols = 2)
	tabela1.style = estilo_tabela
	tabela1.alignment = WD_TABLE_ALIGNMENT.CENTER
	tabela1.rows[0].cells[0].text = 'PROPRIEDADES DO MATERIAL'
	tabela1.rows[1].cells[0].text = u'M�dulo de Elasticidade (MPa)'
	tabela1.rows[2].cells[0].text = 'Coeficiente de Poisson'
	tabela1.rows[3].cells[0].text = u'Tens�o de Escoamento (MPa)'
	tabela1.rows[4].cells[0].text = u'Tens�o de Ruptura (MPa)'
	#tabela1.rows[0].cells[1].text = ''
	tabela1.rows[1].cells[1].text = str(modulo_elasticidade)#str(format(modulo_elasticidade, ".2e"))
	tabela1.rows[2].cells[1].text = str(coeficiente_poisson)#str(format(coeficiente_poisson, ".2f"))
	tabela1.rows[3].cells[1].text = str(tensao_escoamento)#str(format(tensao_escoamento, ".2e"))
	tabela1.rows[4].cells[1].text = str(tensao_ruptura)#str(format(tensao_ruptura, ".2e"))

	tabela1 = configuraTabela(doc, tabela1, 5, 2, titulo_unico = True)

	#GR�FICO
	doc = configuraImagem(doc, arquivo_grafico, legenda = legenda1)

	paragrafo2 = doc.add_paragraph(legenda3)
	paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
	paragrafo2.runs[0].font.name = 'Times New Roman'
	paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)

	#TABELA 2 - TENS�O E DEFORMA��O1
	if tensoes is None:
		tensoes = [0.0, 534.1, 586.356, 607.3325, 631.703, 642.6945, 653.7068, 670.2555, 675.3539, 680.4538, 685.5568, 690.6619, 700.875, 705.984, 718.6498, 721.3181, 729.7026, 740.6435, 754.1836]
	if deformacoes is None:
		deformacoes = [0.0, 0.002671, 0.009, 0.0125, 0.01872, 0.022567, 0.027259, 0.036259, 0.039593, 0.043232, 0.047203, 0.051533, 0.061388, 0.06698, 0.083037, 0.086861, 0.1, 0.12, 0.15]
	tabela2 = doc.add_table(rows = (len(tensoes) + 1), cols = 2)
	tabela2.style = estilo_tabela
	tabela2.alignment = WD_TABLE_ALIGNMENT.CENTER
	tabela2.rows[0].cells[0].text = u'TENS�O (MPa)'
	tabela2.rows[0].cells[1].text = u'DEFORMA��O'
	for i in range(1, len(tensoes) + 1):
		tabela2.rows[i].cells[0].text = str(tensoes[i - 1])#str(format(tensoes[i - 1], ".2e"))
		tabela2.rows[i].cells[1].text = str(deformacoes[i - 1])#str(format(deformacoes[i - 1], ".2e"))

	tabela2 = configuraTabela(doc, tabela2, len(tensoes) + 1, 2, titulo_unico = False)#, legenda = legenda3) 
	
	return doc, num_figura, num_tabela

def modeloDoDefeitoIsolado(doc, num_topico, num_subtopico, num_tabela, num_figura,  profundidade = None, comprimento = None, largura = None, posicao_supeprficie = None, pos_horaria = None, pos_longitudinal = None, arquivo_imagem1 = None, arquivo_imagem2 = None, arquivo_imagem3 = None):

	titulo = str(num_topico) + '.' + str(num_subtopico) + ' MODELO DO DEFEITO ISOLADO'
	texto = [u'As principais caracter�sticas geom�tricas dos defeitos (profundidade ', 'd', ', comprimento ', 'L', ' e largura ', 'W',
                 u', vide Figura 2) a serem analisados est�o descritas na Tabela 5. Da mesma forma, est�o indicados a posi��o do defeito na superf�cie do duto, ',
                 u'bem como os raios de concord�ncia (TR) e ado�amento (FR).']
##	texto = [u'	As principais caracter�sticas geom�tricas dos defeitos (profundidade ',
##                 'D', ', comprimento ', 'LL', ' e largura ', 'LC',
##                 u', vide figura) a serem analisadas est�o descritas na tabela. Al�m do mais, est� indicada a posi��o do defeito na superf�cie do duto.']

	legenda1 = 'Tabela ' + str(num_tabela) + u': Caracter�sticas geom�tricas e posi��es do defeito.'
	num_tabela += 1
	#legenda2 = 'Figura ' + str(num_figura) + ': Detalhe da discretiza��o de um quarto do defeito.'
	legenda3 = 'Figura ' + str(num_figura) + ': Detalhe da geometria do defeito.'
	#legenda4 = 'Figura ' + str(num_figura + 1) + ': Vista da discretiza��o do duto na regi�o do defeito.'
	num_figura += 1

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO
	paragrafo2 = doc.add_paragraph()
	paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	for i in range(len(texto)):
		run = paragrafo2.add_run(texto[i])
		if i in [1, 3, 5]:
			run.font.italic = True
		run.font.name = 'Times New Roman'
		run.font.size = Pt(tamanho_fonte_texto)

	arquivo_imagem_1 = 'esquematico_defeito_retangular.png'
	print('certo_fora')

	doc = configuraImagem(doc, arquivo_imagem_1)#, legenda = u'Detalhe da geometria do defeito de corros�o. Vistas de topo (theta-Z) e no plano longitudinal (R-Z)')
	print('certo_fora_dps')
	#TABELA
	if profundidade is None:
		profundidade = 600
	if comprimento is None:
		comprimento = 10
	if largura is None:
		largura = 5
	if posicao_supeprficie is None:
		posicao_supeprficie = 'Externo'
	if pos_horaria is None:
		pos_horaria = '06:00:00'
	if pos_longitudinal is None:
		pos_longitudinal = '02:00:00'
	tabela = doc.add_table(rows = 8, cols = 2)
	tabela.style = estilo_tabela
	tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
	tabela.rows[0].cells[0].text = 'DADOS DO DEFEITO ISOLADO'
	tabela.rows[1].cells[0].text = 'Profundidade (mm)'
	tabela.rows[2].cells[0].text = 'Comprimento (mm)'
	tabela.rows[3].cells[0].text = 'Largura (mm)'
	tabela.rows[4].cells[0].text = u'Posi��o na Superf�cie do Duto'
	tabela.rows[5].cells[0].text = u'Posi��o Hor�ria do Defeito no Duto (graus)*'
	tabela.rows[6].cells[0].text = u'Posi��o Longitudinal do Defeito no Duto (mm)*'
	#tabela.rows[7].cells[0].text = ''
	#tabela.rows[0].cells[1].text = ''
	tabela.rows[1].cells[1].text = str(format(profundidade, ".2e"))
	tabela.rows[2].cells[1].text = str(format(comprimento, ".2e"))
	tabela.rows[3].cells[1].text = str(format(largura, ".2e"))
	tabela.rows[4].cells[1].text = str(posicao_supeprficie)
	tabela.rows[5].cells[1].text = str(pos_horaria)
	tabela.rows[6].cells[1].text = str(format(pos_longitudinal, ".2e"))
	tabela.rows[7].cells[1].text = '*Medida a partir do canto superior esquerdo do defeito.'

	tabela = configuraTabela(doc, tabela, 8, 2, titulo_unico = True, legenda = legenda1)

	cell70 = tabela.rows[7].cells[0]
	cell71 = tabela.rows[7].cells[1]
	cell70.merge(cell71)
	cell70.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

	#para separar a tabela das imagens
	doc.add_paragraph()

	#FIGURAS
	if arquivo_imagem1 is None:
		arquivo_imagem1 = 'defeito_1.jpg'
	if arquivo_imagem2 is None:
		arquivo_imagem2 = 'DimDefeitoRedimensionada.png'
	if arquivo_imagem3 is None:
		arquivo_imagem3 = 'defeito_isolado.png'
	print('certo_final')
	#doc = configuraImagem(doc, arquivo_imagem1, legenda = legenda2)
	doc = configuraImagem(doc, arquivo_imagem2, legenda = legenda3)
	#doc = configuraImagem(doc, arquivo_imagem3, legenda = legenda4)

	return doc, num_figura, num_tabela

def modeloParaMultiplosDefeitos(doc, num_topico, num_subtopico, num_tabela, num_figura,  qtd_defeitos = None, profundidades = None, comprimentos = None, larguras = None, posicoes_superficie = None, posicoes_horarias = None, posicoes_longitudinais = None, arquivo_imagem1 = None, arquivo_imagem2 = None, arquivo_imagem3 = None):

	titulo = '\n' + str(num_topico) + '.' + str(num_subtopico) + ' MODELO PARA M�LTIPLOS DEFEITOS'
	texto1 = [u'	As principais caracter�sticas geom�tricas dos defeitos (profundidade ', 'D',
                  ', comprimento ', 'LL', ' e largura ', 'LC',
                  u', vide figura) a serem analisadas est�o descritas na tabela. Al�m do mais, est� indicada a posi��o do defeito na superf�cie do duto.']
	texto2 = u'	Na situa��o de m�ltiplos defeitos,  a posi��o hor�ria e posi��o longitudinal dos defeitos s�o informadas na tabela (ambas medidas a partir do canto superior esquerdo do defeito).'

	#legenda2 = 'Figura ' + str(num_figura) + ': Detalhe da discretiza��o de um quarto do defeito.'
	legenda3 = 'Figura ' + str(num_figura) + ': Detalhe da geometria do defeito.'
	#legenda4 = 'Figura ' + str(num_figura + 1) + ': Vista da discretiza��o do duto na regi�o do defeito.'
	num_figura += 1

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO 1
	paragrafo2 = doc.add_paragraph()
	paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
	for i in range(len(texto1)):
		run = paragrafo2.add_run(texto1[i])
		if i in [1, 3, 5]:
			run.font.italic = True
		run.font.name = 'Times New Roman'
		run.font.size = Pt(tamanho_fonte_texto)

	#TEXTO 2
	doc = configuraTexto(doc, texto2)

	#TABELAS
	if qtd_defeitos is None:
		qtd_defeitos = 2
	if profundidades is None:
		profundidades = [600, 300]
	if comprimentos is None:
		comprimentos = [10, 20]
	if larguras is None:
		larguras = [5, 10]
	if posicoes_superficie is None:
		posicoes_superficie = ['Externo', 'Interno']
	if posicoes_horarias is None:
		posicoes_horarias = [10, 15]
	if posicoes_longitudinais is None:
		posicoes_longitudinais = [10, 15]
	for i in range(qtd_defeitos):
		tabela = doc.add_table(rows = 8, cols = 2)
		tabela.style = estilo_tabela
		tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
		tabela.rows[0].cells[0].text = 'DADOS DO DEFEITO ' + str(i + 1)
		tabela.rows[1].cells[0].text = 'Profundidade (mm)'
		tabela.rows[2].cells[0].text = 'Comprimento (mm)'
		tabela.rows[3].cells[0].text = 'Largura (mm)'
		tabela.rows[4].cells[0].text = u'Posi��o na Superf�cie do Duto'
		tabela.rows[5].cells[0].text = u'Posi��o Hor�ria do Defeito no Duto (graus)*'
		tabela.rows[6].cells[0].text = u'Posi��o Longitudinal do Defeito no Duto (mm)*'
		#tabela.rows[7].cells[0].text = ''		
		#tabela.rows[0].cells[1].text = ''
		tabela.rows[1].cells[1].text = str(format(profundidades[i], ".2e"))
		tabela.rows[2].cells[1].text = str(format(comprimentos[i], ".2e"))
		tabela.rows[3].cells[1].text = str(format(larguras[i], ".2e"))
		tabela.rows[4].cells[1].text = str(posicoes_superficie[i])
		tabela.rows[5].cells[1].text = str(posicoes_horarias[i])
		tabela.rows[6].cells[1].text = str(format(posicoes_longitudinais[i], ".2e"))
		tabela.rows[7].cells[1].text = '*Medida a partir do canto superior esquerdo do defeito.'

		legenda = 'Tabela ' + str(num_tabela) + u': Caracter�sticas geom�tricas e posi��es do defeito ' + str(i + 1) + '.'
		num_tabela += 1

		tabela = configuraTabela(doc, tabela, 8, 2, titulo_unico = True, altura_linha = 0.85, legenda = legenda)

		cell70 = tabela.rows[7].cells[0]
		cell71 = tabela.rows[7].cells[1]
		cell70.merge(cell71)
		cell70.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

		doc.add_paragraph()		

	#FIGURAS
	if arquivo_imagem1 is None:
		arquivo_imagem1 = 'defeito_1.jpg'
	if arquivo_imagem2 is None:
		arquivo_imagem2 = 'DimDefeitoRedimensionada.png'
	if arquivo_imagem3 is None:
		arquivo_imagem3 = 'multiplos_defeitos.png'

	#doc = configuraImagem(doc, arquivo_imagem1, legenda = legenda2)
	doc = configuraImagem(doc, arquivo_imagem2, legenda = legenda3)
	#doc = configuraImagem(doc, arquivo_imagem3, legenda = legenda4)

	return doc, num_figura, num_tabela

def discretizacaoDoModeloDeDutoCorroido(doc, num_topico, num_subtopico, num_figura,  arquivo_discr_um_quarto = None, arquivo_discr_prox_defeito = None):

	titulo = str(num_topico) + '.' + str(num_subtopico) + u' DISCRETIZA��O DO MODELO DE DUTO CORRO�DO'
	texto1 = u'	De posse do modelo geom�trico definido, a malha discreta � gerada fazendo-se uso de elementos s�lidos (SOLID45) tridimensionais lineares ao longo de todo modelo.'
	texto2 = u'	A figura abaixo cont�m a discretiza��o de um quarto do duto com um defeito idealizado isolado.'
	legenda1 = 'Figura ' + str(num_figura) + u': Discretiza��o de um quarto do duto com defeito.'
	#legenda2 = 'Figura ' + str(num_figura + 1) + ': Vista da discretiza��o do duto na regi�o pr�xima do defeito.'
	num_figura += 1

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO 1
	doc = configuraTexto(doc, texto1)

	#TEXTO 2
	doc = configuraTexto(doc, texto2)

	#FIGURAS
	if arquivo_discr_um_quarto is None:
		arquivo_discr_um_quarto = 'discretizacao_um_quarto.png'
	if arquivo_discr_prox_defeito is None:
		arquivo_discr_prox_defeito = 'discretizacao_prox_defeito.png'

	doc = configuraImagem(doc, arquivo_discr_um_quarto, legenda = legenda1)
	#doc = configuraImagem(doc, arquivo_discr_prox_defeito, legenda = legenda2)

	return doc, num_figura

def condicoesDeContorno(doc, num_topico, num_subtopico, num_figura,  tipo_de_duto = None, arquivo_restricao_z = None, arquivo_restricao_x = None, arquivo_restricao_y = None, arquivo_pressao_longitudinal = None):

	if tipo_de_duto is None:
		tipo_de_duto = 'restrito'

	#para todos os tipos de dutos, esse texto se repete
	titulo = str(num_topico) + '.' + str(num_subtopico) + u' CONDI��ES DE CONTORNO'
	texto1 = u'	Para este modelo a op��o de duto ' + tipo_de_duto + ' foi adotada.'
	texto2 = u'	As condi��es de contorno aplicadas no duto com defeito, considerando-se a simetria est�o indicadas nas figuras a seguir.'

	#legendas que ir�o variar com o tipo de duto
	legenda = 'Figura ' + str(num_figura) + ': ' + u'Condi��es de contorno, considerando 2 planos de simetria'
	legenda1 = u'Restri��o dos n�s na dire��o Z (simetria e condi��es de extremidade).'
	legenda2 = u'Restri��o dos n�s na dire��o X (movimento de corpo livre e simetria).'
	legenda3 = u'Restri��o dos n�s na dire��o Y (evitar movimento de corpo livre).'
	legenda4 = u'Press�o longitudinal (na extremidade sem defeito).'

	#endere�os das imagens que ir�o variar com o tipo de duto
	if arquivo_restricao_z is None:
		arquivo_restricao_z = 'duto_restrito_c_tampa_z.png'
	if arquivo_restricao_x is None:
		arquivo_restricao_x = 'duto_restrito_c_tampa_x.png'
	if arquivo_restricao_y is None:
		arquivo_restricao_y = 'duto_restrito_c_tampa_y.png'
	if arquivo_pressao_longitudinal is None:
		arquivo_pressao_longitudinal = 'duto_nao_restrito_c_tampa.png'

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO 1
	doc = configuraTexto(doc, texto1)

	#TEXTO 2
	doc = configuraTexto(doc, texto2)

	if tipo_de_duto == 'restrito':
	
		legenda = 'Figura ' + str(num_figura) + ': ' + u'Condi��es de contorno, considerando 2 planos de simetria.'

		background = Image.open("base.png")
		img_z = Image.open("z.png")
		img_x = Image.open("x.png")
		img_y = Image.open("y.png")

		background.paste(img_z, (0, 0), img_z)
		background.paste(img_x, (0, 0), img_x)
		background.paste(img_y, (0, 0), img_y)
		background.save("cond_contorno.png","PNG")

		doc = configuraImagem(doc, "cond_contorno.png", legenda = legenda)

		num_figura += 1

	elif tipo_de_duto == 'n�o restrito':
		
		legenda = 'Figura ' + str(num_figura) + ': ' + u'Condi��es de contorno, considerando 2 planos de simetria.'

		background = Image.open("base.png")
		img_x = Image.open("x.png")
		img_y = Image.open("y.png")
		shell = Image.open("shell.png")

		background.paste(shell, (0, 0), shell)
		background.paste(img_x, (0, 0), img_x)
		background.paste(img_y, (0, 0), img_y)
		background.save("cond_contorno.png","PNG")

		doc = configuraImagem(doc, "cond_contorno.png", legenda = legenda)

		num_figura += 1
	
	elif tipo_de_duto == 'n�o restrito com tampa':
		
		legenda = 'Figura ' + str(num_figura) + ': ' + u'Condi��es de contorno, considerando 2 planos de simetria.'

		background = Image.open("base.png")
		img_x = Image.open("x.png")
		img_y = Image.open("y.png")
		img_pressao = Image.open("pressao_long.png")

		background.paste(img_x, (0, 0), img_x)
		background.paste(img_y, (0, 0), img_y)
		background.paste(img_pressao, (0, 0), img_pressao)
		background.save("cond_contorno.png","PNG")

		doc = configuraImagem(doc, "cond_contorno.png", legenda = legenda)

		num_figura += 1

	return doc, num_figura

def cargasAplicadas(doc, num_topico, num_subtopico, num_figura, pressao_int, pressao_ext, temperatura, momento, forca_axial,  tipo_de_duto = None, arquivo_pressao_interna = None, arquivo_pressao_externa = None, arquivo_momento = None, arquivo_temperatura = None, arquivo_forca_axial = None):

	if tipo_de_duto is None:
		tipo_de_duto = 'restrito'

	titulo = str(num_topico) + '.' + str(num_subtopico) + ' CARGAS APLICADAS'
	texto = u'	Foram aplicadas as seguintes cargas: '
	texto1 = u'	Foi aplicada a seguinte carga: '
	texto2 = u'	Nenhuma carga foi aplicada no caso em an�lise.'

	texto3 = u'uma press�o interna de '
	texto4 = u'uma press�o externa de '
	texto5 = u'uma diferen�a de temperatura de '
	texto6 = u'um momento fletor de '
	texto7 = u'uma for�a axial de '

	texto8 = u' A figura a seguir indica o carregamento que foi aplicado no modelo.'
	texto9 = u' As figuras a seguir indicam os carregamentos que foram aplicados no modelo.'

	#arquivos das imagens
	if arquivo_pressao_interna is None:
		arquivo_pressao_interna = 'pressao_interna_valor_fornecido.png'
	if arquivo_pressao_externa is None:
		arquivo_pressao_externa = 'pressao_externa_valor_fornecido.png'
	if arquivo_momento is None:
		arquivo_momento = 'momento_valor_fornecido.png'
	if arquivo_temperatura is None:
		arquivo_temperatura = 'temperatura.png'
	if arquivo_forca_axial is None:
		arquivo_forca_axial = 'forca_axial.png'

	#legendas das imagens
	legenda1 = u'Press�o interna.'
	legenda2 = u'Press�o externa.'
	legenda3 = u'Momento fletor (�ngulo do momento = 10�).'
	legenda4 = u'Defini��o do valor da press�o fornecido.'
	legenda5 = u'Temperatura (coeficiente de dilata��o linear = 1,2 x 10e-5).'
	legenda6 = u'For�a axial.'
	legenda = u'Carregamento de servi�os.'

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	background = Image.open("base2.png")
	img_p_interna = Image.open("pressao_interna.png")
	img_p_externa = Image.open("pressao_externa.png")
	img_f_axial = Image.open("forca_axial.png")
	img_momento = Image.open("momento.png")

	if tipo_de_duto == 'restrito':

		cargas = [pressao_int, pressao_ext, temperatura]

		qtd_cargas = np.count_nonzero(cargas)

		if qtd_cargas == 0: #n�o foram aplicadas cargas

			doc = configuraTexto(doc, texto2)

		elif qtd_cargas == 1: #apenas uma carga foi aplicada

			legenda_imagem = 'Figura ' + str(num_figura) + ': ' + legenda

			if pressao_int != 0:

				texto_final = texto1 + texto3 + str(format(pressao_int, ".2f")) + 'MPa.' + texto8

				doc = configuraTexto(doc, texto_final)

				background.paste(img_p_interna, (0, 0), img_p_interna)

			if pressao_ext != 0:

				texto_final = texto1 + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.' + texto8

				doc = configuraTexto(doc, texto_final)

				background.paste(img_p_externa, (0, 0), img_p_externa)

			if temperatura != 0:

				texto_final = texto1 + texto5 + str(format(temperatura, ".2f")) + '�C.'

				doc = configuraTexto(doc, texto_final)

			background.save("cargas_aplicadas.png","PNG")
			doc = configuraImagem(doc, "cargas_aplicadas.png", legenda = legenda_imagem)
			num_figura += 1

		else: #duas ou tr�s cargas foram aplicadas
			contador = 0

			#MONTAR O TEXTO
			if pressao_int != 0:

				texto_final = texto + texto3 + str(format(pressao_int, ".2f")) + 'MPa'
				contador += 1

			if pressao_ext != 0:

				if contador == qtd_cargas - 1: #se for a �ltima carga
					texto_final = texto_final + ' e ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.'
				elif contador == 0: #se for a primeira carga
					texto_final = texto + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'
					contador += 1
				else: #se for a carga intermedi�ria
					texto_final = texto_final + ', ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'

			if temperatura != 0:

				texto_final = texto_final + ' e ' + texto5 + str(format(temperatura, ".2f")) + '�C.'

			#TEXTO
			texto_final = texto_final + texto9
			doc = configuraTexto(doc, texto_final)

			#IMAGENS
			legenda_imagem = 'Figura' + str(num_figura) + ': ' + legenda

			if pressao_int != 0:

				background.paste(img_p_interna, (0, 0), img_p_interna)

			if pressao_ext != 0:

				background.paste(img_p_externa, (0, 0), img_p_externa)
			
			background.save("cargas_aplicadas.png","PNG")
			doc = configuraImagem(doc, "cargas_aplicadas.png", legenda = legenda_imagem)
			num_figura += 1

	elif tipo_de_duto == 'n�o restrito':
		
		cargas = [pressao_int, pressao_ext, momento, forca_axial]

		qtd_cargas = np.count_nonzero(cargas)

		if qtd_cargas == 0: #n�o foram aplicadas cargas

			doc = configuraTexto(doc, texto2)

		elif qtd_cargas == 1: #apenas uma carga foi aplicada

			legenda_imagem = 'Figura ' + str(num_figura) + ': ' + legenda

			if pressao_int != 0:

				texto_final = texto1 + texto3 + str(format(pressao_int, ".2f")) + 'MPa.' + texto8

				doc = configuraTexto(doc, texto_final)

				background.paste(img_p_interna, (0, 0), img_p_interna)

			if pressao_ext != 0:

				texto_final = texto1 + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.' + texto8

				doc = configuraTexto(doc, texto_final)

				background.paste(img_p_externa, (0, 0), img_p_externa)

			if momento != 0:

				texto_final = texto1 + texto6 + str(format(momento, ".2f")) + 'kN.m.' + texto8

				doc = configuraTexto(doc, texto_final)

				background.paste(img_momento, (0, 0), img_momento)

			if forca_axial != 0:

				texto_final = texto1 + texto7 + str(format(forca_axial, ".2f")) + 'kN.' + texto8

				doc = configuraTexto(doc, texto_final)

				background.paste(img_f_axial, (0, 0), img_f_axial)

			background.save("cargas_aplicadas.png","PNG")
			doc = configuraImagem(doc, "cargas_aplicadas.png", legenda = legenda_imagem)
			num_figura += 1

		else: #duas, tr�s ou quatro cargas foram aplicadas
			contador = 0

			legenda_imagem = 'Figura ' + str(num_figura) + ': ' + legenda

			#MONTAR O TEXTO
			if pressao_int != 0:

				texto_final = texto + texto3 + str(format(pressao_int, ".2f")) + 'MPa'
				contador += 1

			if pressao_ext != 0:

				if contador == qtd_cargas - 1: #se for a �ltima carga
					texto_final = texto_final + ' e ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.'
				elif contador == 0: #se for a primeira carga
					texto_final = texto + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'
					contador += 1
				else: #se for a carga intermedi�ria
					texto_final = texto_final + ', ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'
					contador += 1

			if momento != 0:

				if contador == qtd_cargas - 1: #se for a �ltima carga
					texto_final = texto_final + ' e ' + texto6 + str(format(momento, ".2f")) + 'kN.m.'
				elif contador == 0: #se for a primeira carga
					texto_final = texto + texto6 + str(format(momento, ".2f")) + 'kN.m'
					contador += 1
				else: #se for a carga intermedi�ria
					texto_final = texto_final + ', ' + texto6 + str(format(momento, ".2f")) + 'kN.m'
					contador += 1

			if forca_axial != 0:

				texto_final = texto_final + ' e ' + texto7 + str(format(forca_axial, ".2f")) + 'kN.'

			#TEXTO
			texto_final = texto_final + texto9
			doc = configuraTexto(doc, texto_final)

			#IMAGENS
			if pressao_int != 0:

				background.paste(img_p_interna, (0, 0), img_p_interna)

			if pressao_ext != 0:

				background.paste(img_p_externa, (0, 0), img_p_externa)

			if momento != 0:

				background.paste(img_momento, (0, 0), img_momento)

			if forca_axial != 0:

				background.paste(img_f_axial, (0, 0), img_f_axial)

			background.save("cargas_aplicadas.png","PNG")
			doc = configuraImagem(doc, "cargas_aplicadas.png", legenda = legenda_imagem)
			num_figura += 1

	elif tipo_de_duto == 'n�o restrito com tampa':
		
		cargas = [pressao_int, pressao_ext, forca_axial]

		qtd_cargas = np.count_nonzero(cargas)

		if qtd_cargas == 0: #n�o foram aplicadas cargas

			doc = configuraTexto(doc, texto2)

		elif qtd_cargas == 1: #apenas uma carga foi aplicada

			legenda_imagem = 'Figura ' + str(num_figura) + ': ' + legenda

			if pressao_int != 0:

				texto_final = texto1 + texto3 + str(format(pressao_int, ".2f")) + 'MPa.' + texto8

				doc = configuraTexto(doc, texto_final)

				background.paste(img_p_interna, (0, 0), img_p_interna)

			if pressao_ext != 0:

				texto_final = texto1 + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.' + texto8

				doc = configuraTexto(doc, texto_final)

				background.paste(img_p_externa, (0, 0), img_p_externa)

			if forca_axial != 0:

				texto_final = texto1 + texto7 + str(format(forca_axial, ".2f")) + 'kN.' + texto8

				doc = configuraTexto(doc, texto_final)

				background.paste(img_f_axial, (0, 0), img_f_axial)

			background.save("cargas_aplicadas.png","PNG")
			doc = configuraImagem(doc, "cargas_aplicadas.png", legenda = legenda_imagem)
			num_figura += 1

		else: #duas ou tr�s cargas foram aplicadas

			contador = 0

			legenda_imagem = 'Figura ' + str(num_figura) + ': ' + legenda

			#MONTAR O TEXTO
			if pressao_int != 0:

				texto_final = texto + texto3 + str(format(pressao_int, ".2f")) + 'MPa'
				contador += 1

			if pressao_ext != 0:

				if contador == qtd_cargas - 1: #se for a �ltima carga
					texto_final = texto_final + ' e ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.'
				elif contador == 0: #se for a primeira carga
					texto_final = texto + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'
					contador += 1
				else: #se for a carga intermedi�ria
					texto_final = texto_final + ', ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'

			if forca_axial != 0:

				texto_final = texto_final + ' e ' + texto7 + str(format(forca_axial, ".2f")) + 'kN.'

			#TEXTO
			texto_final = texto_final + texto9
			doc = configuraTexto(doc, texto_final)

			#IMAGENS
			if pressao_int != 0:

				background.paste(img_p_interna, (0, 0), img_p_interna)

			if pressao_ext != 0:

				background.paste(img_p_externa, (0, 0), img_p_externa)

			if forca_axial != 0:

				background.paste(img_f_axial, (0, 0), img_f_axial)

			background.save("cargas_aplicadas.png","PNG")
			doc = configuraImagem(doc, "cargas_aplicadas.png", legenda = legenda_imagem)
			num_figura += 1

	return doc, num_figura

def analisePressaoFornecida(doc, num_topico, num_subtopico,  pressao_fornecida = None, arquivo_pressao_fornecida = None):

	if pressao_fornecida is None:
		pressao_fornecida = 10
	if arquivo_pressao_fornecida is None:
		arquivo_pressao_fornecida = 'pressao_fornecida.png'

	titulo = str(num_topico) + '.' + str(num_subtopico) + u' AN�LISE AT� O VALOR DE PRESS�O FORNECIDO'
	texto = u'	A figura abaixo indica que o sistema tentar� executar a an�lise at� o valor da press�o fornecida de ' + str(format(pressao_fornecida, ".2f")) + 'MPa.'
	texto1 = u'	O sistema tentar� executar a an�lise at� o valor da press�o fornecida de ' + str(format(pressao_fornecida, ".2f")) + 'MPa.'

	#legenda = 'Figura ' + str(num_figura) + ': Defini��o do valor da press�o fornecido.'

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO 
	doc = configuraTexto(doc, texto1)

	#FIGURA
	#doc = configuraImagem(doc, arquivo_pressao_fornecida, legenda = legenda)

	return doc

def analisePressaoRuptura(doc, num_topico, num_subtopico,  arquivo_pressao_ruptura = None):

	if arquivo_pressao_ruptura is None:
		arquivo_pressao_ruptura = 'pressao_ruptura.png'

	titulo = str(num_topico) + '.' + str(num_subtopico) + u' AN�LISE AT� VALOR DE PRESS�O DE RUPTURA'
	texto = u'	A figura abaixo indica que o sistema tentar� executar a an�lise at� o valor da press�o de ruptura.'
	texto1 = u'	O sistema tentar� executar a an�lise at� o valor da press�o de ruptura.'

	#legenda = 'Figura ' + str(num_figura) + ': Escolha para an�lise de ruptura.'

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO 
	doc = configuraTexto(doc, texto1)

	#FIGURA
	#doc = configuraImagem(doc, arquivo_pressao_ruptura, legenda = legenda)

	return doc

def resultadosEscoamento(doc, num_topico, num_subtopico, num_figura,  arquivo_mapa_tensoes = None, arquivo_janela_CORDUT = None):

	titulo = str(num_topico) + '.' + str(num_subtopico) + u' RESULTADOS � N�VEL DE ESCOAMENTO'
	texto = u'	Na figura abaixo est� apresentado o detalhe do mapa de tens�es do modelo no passo de carga em que se iniciou o escoamento.'

	#arquivos das imagens
	if arquivo_mapa_tensoes is None:
		arquivo_mapa_tensoes = 'mapa_tensoes.png'
	if arquivo_janela_CORDUT is None:
		arquivo_janela_CORDUT = 'janela_CORDUT.png'

	#legendas das imagens
	legenda1 = 'Figura ' + str(num_figura) + u': Tens�es do modelo no passo de carga em que se iniciou o escoamento.'
	#legenda2 = 'Figura ' + str(num_figura + 1) + ': Janela de resultado do CORDUT no in�cio do escoamento.'

	num_figura += 1

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO 1
	doc = configuraTexto(doc, texto)

	#FIGURA 1
	doc = configuraImagem(doc, arquivo_mapa_tensoes, legenda = legenda1)

	#FIGURA 2
	#doc = configuraImagem(doc, arquivo_janela_CORDUT, legenda = legenda2)	

	return doc, num_figura

def resultadosPressaoFornecida(doc, num_topico, num_subtopico, num_figura,  arquivo_tensao_pressao = None, arquivo_janela_resultado = None):

	titulo = str(num_topico) + '.' + str(num_subtopico) + u' RESULTADO � N�VEL DE VALOR DE PRESS�O FORNECIDO'
	texto = u'	Na figura abaixo est� apresentada a perspectiva do mapa de tens�es do modelo do duto com defeito no passo de carga que atingiu a press�o fornecida.'

	#legendas das imagens
	legenda1 = 'Figura ' + str(num_figura) + u': Resultado da an�lise: tens�o e press�o no passo de carga que atingiu a press�o fornecida.'
	#legenda2 = 'Figura ' + str(num_figura + 1) + ': Janela de resultado do CORDUT no passo de carga que atingiu a press�o fornecida.'
	num_figura += 1

	#arquivo das imagens
	if arquivo_tensao_pressao is None:
		arquivo_tensao_pressao = 'tensao_pressao.png'
	if arquivo_janela_resultado is None:
		arquivo_janela_resultado = 'janela_CORDUT_resultado.png'

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO 1
	doc = configuraTexto(doc, texto)

	#FIGURA 1
	doc = configuraImagem(doc, arquivo_tensao_pressao, legenda = legenda1)

	#FIGURA 2
	#doc = configuraImagem(doc, arquivo_janela_resultado, legenda = legenda2)

	return doc, num_figura

def resultadosPressaoRuptura(doc, num_topico, num_subtopico, num_figura,  arquivo_tensao_pressao = None, arquivo_janela_resultado = None):

	titulo = str(num_topico) + '.' + str(num_subtopico) + u' RESULTADO � N�VEL DE RUPTURA'
	texto = u'	Na figura abaixo est� apresentada a perspectiva do mapa de tens�es do modelo do duto com defeito na ruptura.'

	#legenda das imagens
	legenda1 = 'Figura ' + str(num_figura) + ': Resultado da an�lise: tens�o m�xima e press�o �ltima.'
	#legenda2 = 'Figura ' + str(num_figura + 1) + ': Janela de resultado do CORDUT no final da an�lise.'
	num_figura += 1

	#arquivo das imagens
	if arquivo_tensao_pressao is None:
		arquivo_tensao_pressao = 'tensao_pressao_ruptura.png'
	if arquivo_janela_resultado is None:
		arquivo_janela_resultado = 'janela_CORDUT_resultado_ruptura.png'

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO 1
	doc = configuraTexto(doc, texto)

	#FIGURA 1
	doc = configuraImagem(doc, arquivo_tensao_pressao, legenda = legenda1)

	#FIGURA 2
	#doc = configuraImagem(doc, arquivo_janela_resultado, legenda = legenda2)

	return doc, num_figura

def deslocamentos(doc, num_topico, num_subtopico, num_figura,  analise_ruptura = None, arquivo_deslocamento = None):

	if analise_ruptura is None:
		analise_ruptura = False

	titulo = str(num_topico) + '.' + str(num_subtopico) + ' DESLOCAMENTOS'
	texto1 = u'	A deformada do duto na regi�o do defeito (em escala autom�tica do ANSYS), referente ao valor '
	texto2 = u'de press�o fornecida'
	texto3 = u'da press�o de ruptura'
	texto4 = u', est� ilustrada na figura a seguir. Como esperado, deslocamentos m�ximos ocorreram na regi�o do defeito.'
	legenda1 = 'Figura ' + str(num_figura) + u': Configura��o deformada do modelo do duto com defeito referente ao valor '
	num_figura += 1
	legenda2 = '.'
	if analise_ruptura:
		texto = texto1 + texto3 + texto4
		legenda = legenda1 + texto3 + legenda2
	else:
		texto = texto1 + texto2 + texto4
		legenda = legenda1 + texto2 + legenda2

	#arquivo da imagem
	if arquivo_deslocamento is None:
		arquivo_deslocamento = 'deslocamento.png'

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO
	doc = configuraTexto(doc, texto)

	#FIGURA
	doc = configuraImagem(doc, arquivo_deslocamento, legenda = legenda)

	return doc, num_figura

def recomendacoes(doc, num_topico,  recomendacao_usuario = None):

	titulo = str(num_topico) + u'. RECOMENDA��ES'

	default1 = [u'	A fim de avaliar o comprometimento da integridade estrutural dos casos dos dutos analisados, os resultados da press�o de um duto �ntegro (especificar valor) e com defeito indicam que o duto poder� ser mantido em opera��o com carga inferior � de ruptura calculada (especificar valor).',
                    u'	Entretanto, recomenda-se definir a seguran�a da estrutura atrav�s do emprego de um fator de seguran�a fornecido pelas normas de projeto ou at� mesmo atrav�s da an�lise de confiabilidade. Esta permite que inspe��es e reparos sejam programados de forma a garantir a seguran�a do duto corro�do.']
	default2 = [u'	O CORDUT concluiu a simula��o com sucesso.', u'	A tens�o de escoamento corresponde a um valor de 534 Mpa.',
                    u'	A tens�o referente a um valor de press�o fornecida corresponde a um valor de 620 Mpa.',
                    u'	Neste caso, a ruptura n�o foi detectada (tens�o de ruptura 754,18 Mpa).', u'	O duto poder� manter em opera��o com seguran�a.']

	#T�TULO
	doc = configuraTexto(doc, titulo, titulo = True)

	#TEXTO
	'''
	if recomendacao_usuario is None:
		for frase in default1:
			doc = configuraTexto(doc, frase)
	else:
		doc = configuraTexto(doc, recomendacao_usuario)
	'''
	doc = configuraTexto(doc, '	')

	return doc 

def gerarRelatorio(dicionario):

	#organizadores gerais do relat�rio
	num_topico = 1
	num_tabela = 1
	num_figura = 1

	#cria��o do documento
	doc = docx.Document()

        #inicializa��o dos dados de identifica��o do relat�rio
	doc = identificacaoDoRelatorio(doc, num_topico)

	#inicializa��o da quantidade de casos analisados no relat�rio
	num_topico += 1
	qtd_casos = len(dicionario)
	doc, num_tabela = quantidadeDeCasos(doc, num_topico, num_tabela, qtd_casos = qtd_casos)

	for i in range(qtd_casos):

		#obtendo as informa��es do caso a ser analisado
		caso = dicionario[i + 1]
		
		#obtendo as informa��es das imagens que ir�o compor o relat�rio
		imagens = caso['imagens']

		#obtendo as informa��es do duto
		informacoes_duto = caso['duto']

		#organizador de cada caso analisado do relat�rio
		num_subtopico = 1
		num_topico += 1
		doc, num_tabela = modeloDoDuto(doc, num_topico, num_subtopico, num_tabela, diametro_externo = informacoes_duto['diam_ext'], espessura = informacoes_duto['espessura'], comprimento = informacoes_duto['comp_duto'], taxa_ovalizacao = informacoes_duto['ovalizacao'], nome_modelo = caso['nome_modelo'])

		#obtendo as informa��es do material
		informacoes_material = caso['material']

		#obtendo as informa��es da curva tens�o x deforma��o do material
		duplas_tens_def = informacoes_material['lista_curva_material']
		lista_tensoes = []
		lista_deformacoes = []
		for dupla in duplas_tens_def:
			lista_tensoes.append(dupla[0])
			lista_deformacoes.append(dupla[1])

		#propriedades do material
		num_subtopico += 1
		doc, num_figura, num_tabela = propriedadesDoMaterial(doc, num_topico, num_subtopico, num_tabela, num_figura, arquivo_grafico = informacoes_material['dir_arq_cf_plot'], modulo_elasticidade = informacoes_material['mod_elast'], coeficiente_poisson = informacoes_material['coef_poisson'], tensao_escoamento = informacoes_material['tens_esc'], tensao_ruptura = informacoes_material['tens_rup'], tensoes = lista_tensoes, deformacoes = lista_deformacoes)

		#obtendo as informa��es sobre os defeitos
		informacoes_defeitos = caso['defeitos']
		qtd_defeitos = len(informacoes_defeitos)

		if qtd_defeitos == 1: #DEFEITO ISOLADO

			#obtendo as informa��es sobre o defeito isolado
			defeito = informacoes_defeitos[1]

			num_subtopico += 1
			doc, num_figura, num_tabela = modeloDoDefeitoIsolado(doc, num_topico, num_subtopico, num_tabela, num_figura, profundidade = defeito['prof'], comprimento = defeito['comp_def'], largura = defeito['larg_def'], posicao_supeprficie = defeito['superficie'].capitalize(), pos_horaria = defeito['pos_hor'], pos_longitudinal = defeito['pos_long'])

		else: #M�LTIPLOS DEFEITOS

			#organizando as informa��es dos defeitos
			lista_profundidades = []
			lista_comprimentos = []
			lista_larguras = []
			lista_pos_superficie = []
			lista_pos_hor = []
			lista_pos_long = []
			for i in range(qtd_defeitos):
				defeito = informacoes_defeitos[i + 1]

				lista_profundidades.append(defeito['prof'])
				lista_comprimentos.append(defeito['comp_def'])
				lista_larguras.append(defeito['larg_def'])
				lista_pos_superficie.append(defeito['superficie'].capitalize())
				lista_pos_hor.append(defeito['pos_hor'])
				lista_pos_long.append(defeito['pos_long'])

			num_subtopico += 1
			doc, num_figura, num_tabela = modeloParaMultiplosDefeitos(doc, num_topico, num_subtopico, num_tabela, num_figura, qtd_defeitos = qtd_defeitos, profundidades = lista_profundidades, comprimentos = lista_comprimentos, larguras = lista_larguras, posicoes_superficie = lista_pos_superficie, posicoes_horarias = lista_pos_hor, posicoes_longitudinais = lista_pos_long)
			
		num_subtopico += 1
		doc, num_figura = discretizacaoDoModeloDeDutoCorroido(doc, num_topico, num_subtopico, num_figura, arquivo_discr_um_quarto = imagens['disc_um_quarto_duto'])
		
		num_subtopico += 1
		doc, num_figura = condicoesDeContorno(doc, num_topico, num_subtopico, num_figura, tipo_de_duto = caso['cond_ext'])

		#obtendo as informa��es sobre as cargas aplicadas
		cargas_aplicadas = caso['carreg_serv']
		num_subtopico += 1
		doc, num_figura = cargasAplicadas(doc, num_topico, num_subtopico, num_figura, pressao_int = cargas_aplicadas['pint'], pressao_ext = cargas_aplicadas['pext'], temperatura = cargas_aplicadas['temp'], momento = cargas_aplicadas['momento'], forca_axial = cargas_aplicadas['faxial'], tipo_de_duto = caso['cond_ext'])

		#obtendo as informa��es sobre a an�lise
		analise = caso['analise']

		tipo_analise = analise['condicao_parada']
		if tipo_analise == 'press_fornecida':

			num_subtopico += 1
			doc = analisePressaoFornecida(doc, num_topico, num_subtopico, pressao_fornecida = analise['pressao_solicitante'])

			#ter� escoamento se a press�o fornecida for maior ou igual a press�o de escoamento
			if analise['pressao_solicitante'] >= analise['press_esc']:
				num_subtopico += 1
				doc, num_figura = resultadosEscoamento(doc, num_topico, num_subtopico, num_figura, arquivo_mapa_tensoes = imagens['mapa_tensoes_inicio_escoamento'])
				
			num_subtopico += 1
			doc, num_figura = resultadosPressaoFornecida(doc, num_topico, num_subtopico, num_figura, arquivo_tensao_pressao = imagens['res_analise'])
			
			num_subtopico += 1
			doc, num_figura = deslocamentos(doc, num_topico, num_subtopico, num_figura, analise_ruptura = False, arquivo_deslocamento = imagens['deslocamentos'])

		else:

			num_subtopico += 1
			doc = analisePressaoRuptura(doc, num_topico, num_subtopico)

			#quando h� an�lise at� ruptura, h� resultado � n�vel de escoamento
			num_subtopico += 1
			doc, num_figura = resultadosEscoamento(doc, num_topico, num_subtopico, num_figura, arquivo_mapa_tensoes = imagens['mapa_tensoes_inicio_escoamento'], arquivo_janela_CORDUT = imagens['janela_res_inicio_escoamento'])
			
			num_subtopico += 1
			doc, num_figura = resultadosPressaoRuptura(doc, num_topico, num_subtopico, num_figura, arquivo_tensao_pressao = imagens['res_analise'])
			
			num_subtopico += 1
			doc, num_figura = deslocamentos(doc, num_topico, num_subtopico, num_figura, analise_ruptura = True, arquivo_deslocamento = imagens['deslocamentos'])

	num_topico += 1
	doc = recomendacoes(doc, num_topico)

	#salvando o relat�rio
	doc.save('relatorio.docx')

	#abrindo o arquivo automaticamente
	#os.system('start relatorio.docx')

	return doc

##########           TESTE            ############
dicionario = {1: {'nome_modelo': 'idts_013', 'duto': {'diam_ext': 458.8, 'espessura': 8.2,
                                                      'comp_duto': 2340, 'ovalizacao': 2},
                  'material': {'dir_arq_cf_plot': 'tensao_versus_deformacao.png',
                               'mod_elast': 199962.5608386372, 'coef_poisson': 0.3,
                               'tens_esc': 534.1, 'tens_rup':754.1836,
                               'lista_curva_material': [(0.0, 0.0),(534.1, 0.002671),
                                                        (586.356, 0.009), (607.3325, 0.0125),
                                                        (631.703, 0.0187197), (642.6945, 0.0225671),
                                                        (653.7068, 0.0272592), (670.2555, 0.0362591),
                                                        (675.3539, 0.039593), (680.4538, 0.0432324),
                                                        (685.5568, 0.0472031), (690.6619, 0.0515325),
                                                        (700.875, 0.061388), (705.984, 0.0669801),
                                                        (718.6498, 0.0830373), (721.3181, 0.0868612),
                                                        (729.7026, 0.1), (740.6435, 0.12),
                                                        (754.1836, 0.15)]}, 'defeitos':
                  {1: {'tipo_def': 'retangular', 'comp_def': 50, 'larg_def': 50, 'prof': 4,
                       'superficie': 'interna', 'pos_long': 1145, 'pos_hor': '06:00:00'}},
                  'cond_ext': 'restrito', 'carreg_serv': {'pint': 10.0, 'pext': 2.0,
                                                          'temp': 0.0, 'alfa': 0.0,
                                                          'momento': 0.0, 'faxial': 0.0,
                                                          'ang_momento': 0.0},
                  'analise': {'press_parada': 15, 'tensao_parada': 600, 'press_esc': 10,
                              'tensao_esc': 534.1, 'condicao_parada': 'press_fornecida',
                              'pressao_solicitante': 15.0},
                  'imagens': {'disc_um_quarto_duto': 'discretizacao_um_quarto.png',
                              'mapa_tensoes_inicio_escoamento': 'mapa_tensoes.png',
                              'res_analise': 'tensao_pressao.png', 'deslocamentos': 'deslocamento.png'}}}

gerarRelatorio(dicionario)
