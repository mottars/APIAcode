#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Feb  7 16:15:16 2023

@author: pimentel
"""

#imports necessários
import os
import docx
import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt
from PIL import Image
import math

# from docx2pdf import convert

#parâmetros globais
tamanho_fonte_titulo = 15
tamanho_fonte_texto = 12
tamanho_fonte_tabela = 10
tamanho_fonte_legenda = 10
estilo_tabela = 'Medium Grid 1 Accent 1'
estilo_tabela = 'Medium Shading 1 Accent 1'
estilo_tabela = 'Light Grid Accent 1'

#funções auxiliares
def configuraTabela(doc, tabela, linhas, colunas, titulo_unico,  altura_linha = None, legenda = None, negrito=None):
    if altura_linha is None:
        altura_linha = 0.6
    for i in range(0, linhas):
        tabela.rows[i].height = Cm(altura_linha)
        for j in range(0, colunas):
            if tabela.rows[i].cells[j].text != '':
                tabela.rows[i].cells[j].width = Cm(7)
                tabela.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    tabela.rows[i].cells[j].paragraphs[0].runs[0].font.name = 'Times New Roman'
                    if negrito == False and i != 0:
                                                tabela.rows[i].cells[j].paragraphs[0].runs[0].font.bold = False
                except:
                    print('erro')
                try:
                    tabela.rows[i].cells[j].paragraphs[0].runs[0].font.size = Pt(tamanho_fonte_tabela)
                except:
                    print('erro 2')
                    
                tabela.rows[i].cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if titulo_unico:
        cell00 = tabela.columns[0].cells[0]
        cell01 = tabela.columns[1].cells[0]
        cell00.merge(cell01)
        
    if legenda is not None:
        paragrafo1 = doc.add_paragraph(legenda)
        paragrafo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafo1.runs[0].font.name = 'Times New Roman'
        paragrafo1.runs[0].font.size = Pt(tamanho_fonte_legenda)

    return tabela    

def configuraImagem(doc, arquivo,  legenda = None, flag_tam = None):

    paragrafo = doc.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_imagem = paragrafo.add_run()
    # if flag_tam != None:
    #     run_imagem.add_picture(arquivo, width=Cm(16), height=Cm(12))
    # else:
    #     run_imagem.add_picture(arquivo)
        
    print('certo legenda')
    if legenda is not None:
        print('certo legenda_2')
        paragrafo1 = doc.add_paragraph(legenda)
        paragrafo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafo1.runs[0].font.name = 'Times New Roman'
        print('certo legenda_3')
        paragrafo1.runs[0].font.size = Pt(tamanho_fonte_legenda)
        
    print('certo saiu')
    return doc

def configuraTexto(doc, texto,  titulo = False):

    if titulo:
        paragrafo = doc.add_paragraph(texto)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragrafo.runs[0].font.name = 'Times New Roman'
        paragrafo.runs[0].font.size = Pt(tamanho_fonte_titulo)
        paragrafo.runs[0].font.bold = True
    else:
        paragrafo = doc.add_paragraph(texto)
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragrafo.runs[0].font.name = 'Times New Roman'
        paragrafo.runs[0].font.size = Pt(tamanho_fonte_texto)

    return doc

#funções construtoras do documento
#quantidadeDeCasos e identificacaoDoRelatorio são chamadas apenas uma vez!
#o restante das funções são chamadas uma vez para cada caso que for compor o relatório!
def quantidadeDeCasos(doc, num_topico, num_tabela,  qtd_casos = None):

    titulo = str(num_topico) + u'. QUANTIDADE DE CASOS'
    texto = u'    Neste relatório, será analisada a quantidade de casos de defeitos de corrosão que consta especificada na Tabela 1 a seguir.'
    texto1 = u'QUANTIDADE DE CASOS'
    legenda = u'Tabela ' + str(num_tabela) + u': Quantidade de casos analisados'
    num_tabela += 1

    if qtd_casos is None:
        qtd_casos = 1

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO
    doc = configuraTexto(doc, texto)

    paragrafo1 = doc.add_paragraph(legenda)
    paragrafo1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo1.runs[0].font.name = 'Times New Roman'
    paragrafo1.runs[0].font.size = Pt(tamanho_fonte_legenda)
    
    #TABELA
    tabela = doc.add_table(rows = 1, cols = 2)
    tabela.style = estilo_tabela
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    linha_tabela = tabela.rows[0].cells #primeira linha da tabela como uma lista
    tabela.rows[0].height = Cm(0.6)
    linha_tabela[0].text = texto1 #elemento (0, 0) da tabela
    linha_tabela[1].text = str(qtd_casos) #elemento (0, 1) da tabela
    for item in linha_tabela:
        item.width = Cm(5)
        item.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        item.paragraphs[0].runs[0].font.name = 'Times New Roman'
        item.paragraphs[0].runs[0].font.size = Pt(tamanho_fonte_tabela)
        item.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
        doc.add_paragraph()
    

    return doc, num_tabela

def identificacaoDoRelatorio(doc, num_topico,  relatorio = None, projeto = None, objetivo = None):

    if relatorio is None:
        relatorio = u'Análise de Dutos com Defeitos Isolados.'
        relatorio = ' '
    if projeto is None:
        projeto = u'Avaliação de Integridade Estrutural de Dutos Terrestres.'
        projeto = ' '
    if objetivo is None:
        objetivo = u'Este relatório tem como principal objetivo apresentar os resultados das análises de dutos com defeitos isolados, obtidos com o sistema CORDUT.'
        objetivo = ' '

    titulo = str(num_topico) + u'. Relatório de Avaliação de Duto & Manual de Manutenção Futura'

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #RELATÓRIO
    paragrafo2 = doc.add_paragraph(u'Duto: xxxx')
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo2.runs[0].font.bold = True
    paragrafo2.add_run(relatorio)
    for run in paragrafo2.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    #PROJETO
    paragrafo3 = doc.add_paragraph('Trecho: xxxx')
    paragrafo3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo3.runs[0].font.bold = True
    paragrafo3.add_run(projeto)
    for run in paragrafo3.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    #OBJETIVO
    paragrafo4 = doc.add_paragraph('Engenheiro Responsável: xxxx\n')
    paragrafo4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragrafo4.runs[0].font.bold = True
    paragrafo4.add_run(objetivo)
    for run in paragrafo4.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    return doc

def modeloDoDuto(doc, num_topico, num_subtopico, num_tabela,  diametro_externo = None, espessura = None, comprimento = None, taxa_ovalizacao = None, nome_modelo = None):

    titulo = str(num_topico) + '. CASO ' + str(num_topico - 2)

    titulo1 = str(num_topico) + '.' + str(num_subtopico) + ' MODELO DO DUTO'
    texto = [u'    As principais características geométricas dos segmentos de duto (diâmetro Externo ', 'D', 'e', ', comprimento ', 'L', 'duto',
                 ', espessura de parede ', 't',
                 u') analisados estão descritas na Tabela 2. Além disso, características referente à ovalização inicial da seção transversal do segmento '
                 u'de duto onde se encontra o defeito de corrosão e ao nome do modelo analisado também estão contidas nesta tabela.']

    legenda = 'Tabela ' + str(num_tabela) + u': Características geométricas do duto.'
    num_tabela += 1

    #TÍTULOS
    doc = configuraTexto(doc, titulo, titulo = True)
    doc = configuraTexto(doc, titulo1, titulo = True)

    #TEXTO
    paragrafo2 = doc.add_paragraph()
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i in range(len(texto)):
        run = paragrafo2.add_run(texto[i])
        if i in [1, 2, 4, 5, 7]:
            run.font.italic = True
            if i == 2 or i== 5:
                run.font.subscript = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    #TABELA
    if diametro_externo is None:
        diametro_externo = 600
    if espessura is None:
        espessura = 10
    if comprimento is None:
        comprimento = 3000
    if taxa_ovalizacao is None:
        taxa_ovalizacao = 5
    if nome_modelo is None:
        nome_modelo = 'DT1'

    paragrafo2 = doc.add_paragraph(legenda)
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo2.runs[0].font.name = 'Times New Roman'
    paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)
    
    tabela = doc.add_table(rows = 6, cols = 2)
    tabela.style = estilo_tabela
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.rows[0].cells[0].text = u'DADOS DO SEGMENTO DO DUTO'
##    tabela.rows[1].cells[0].text = u'Diâmetro Externo: De (mm)'
    lista_de_tab = [u'Diâmetro Externo (','D', 'e', ') (mm)']
    paragrafo_celula = tabela.rows[1].cells[0].add_paragraph()
    paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(len(lista_de_tab)):
        run = paragrafo_celula.add_run(lista_de_tab[i])
        if i in [1, 2]:
            run.font.italic = True
            if i == 2:
                run.font.subscript = True
        run.font.name = 'Times New Roman'
##    tabela.rows[2].cells[0].text = u'Espessura do Duto t (mm)'
    lista_de_tab = [u'Espessura do Duto (','t', ') (mm)']
    paragrafo_celula = tabela.rows[2].cells[0].add_paragraph()
    paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(len(lista_de_tab)):
        run = paragrafo_celula.add_run(lista_de_tab[i])
        if i in [1]:
            run.font.italic = True
##            if i == 2:
##                run.font.subscript = True
        run.font.name = 'Times New Roman'

##    tabela.rows[3].cells[0].text = u'Comprimento do Duto Lduto (mm)'
    lista_de_tab = [u'Comprimento do Duto (','L', 'duto', ') (mm)']
    paragrafo_celula = tabela.rows[3].cells[0].add_paragraph()
    paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(len(lista_de_tab)):
        run = paragrafo_celula.add_run(lista_de_tab[i])
        if i in [1, 2]:
            run.font.italic = True
            if i == 2:
                run.font.subscript = True
        run.font.name = 'Times New Roman'
##    tabela.rows[4].cells[0].text = u'Taxa de Ovalização fo (%)'
    lista_de_tab = [u'Taxa de Ovalização (','f', 'o', ') (%)']
    paragrafo_celula = tabela.rows[4].cells[0].add_paragraph()
    paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(len(lista_de_tab)):
        run = paragrafo_celula.add_run(lista_de_tab[i])
        if i in [1, 2]:
            run.font.italic = True
            if i == 2:
                run.font.subscript = True  
        run.font.name = 'Times New Roman'

    tabela.rows[5].cells[0].text = u'Nome do Modelo'
    #tabela.rows[0].cells[1].text = ''
    tabela.rows[1].cells[1].text = str(diametro_externo)#str(format(diametro_externo, ".2e"))
    tabela.rows[2].cells[1].text = str(espessura)#str(format(espessura, ".2e"))
    tabela.rows[3].cells[1].text = str(comprimento)#str(format(comprimento, ".2e"))
    tabela.rows[4].cells[1].text = str(taxa_ovalizacao*100)#str(format(taxa_ovalizacao, ".2f"))
    tabela.rows[5].cells[1].text = nome_modelo

    tabela = configuraTabela(doc, tabela, 6, 2, titulo_unico = True)#, legenda = legenda)
    doc.add_paragraph()
    return doc, num_tabela

def propriedadesDoMaterial(doc, num_topico, num_subtopico, num_tabela, num_figura,  arquivo_grafico = None, modulo_elasticidade = None, coeficiente_poisson = None, tensao_escoamento = None, tensao_ruptura = None, tensoes = None, deformacoes = None):

    titulo = str(num_topico) + '.' + str(num_subtopico) + ' PROPRIEDADES DO MATERIAL'
    texto = [u'    As propriedades do material do duto encontram-se na Tabela 3. A partir desses dados, foi construída a curva tensão verdadeira ', 'versus',
                 u' deformação total verdadeira do material utilizada nas análises, indicada na Figura ' + str(num_figura) + u' e os seus pontos estão apresentados na Tabela 4.']
##                A mesma resulta da combinação de um trecho linear da origem (0,0) até a tensão de escoamento, mais um trecho não linear definido segundo a equação de Ramberg-Osgood.']
    if arquivo_grafico is None:
        arquivo_grafico = os.path.join(os.getcwd(), 'GeradorRelatorio', 'tensao_versus_deformacao.png')

    legenda1 = 'Figura ' + str(num_figura) + u': Curva tensão - deformação verdadeira do material'
    num_figura += 1
    legenda2 = 'Tabela ' + str(num_tabela) + ': Propriedades do material do duto'
    legenda3 = 'Tabela ' + str(num_tabela + 1) + u': Pontos da curva tensão verdadeira versus deformação verdadeira do material'
    num_tabela += 2

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO
    paragrafo2 = doc.add_paragraph()
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i in range(len(texto)):
        run = paragrafo2.add_run(texto[i])
        if i == 1:
            run.font.italic = True
            run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    paragrafo2 = doc.add_paragraph(legenda2)
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo2.runs[0].font.name = 'Times New Roman'
    paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)
    
    #TABELA 1 - PROPRIEDADES DO MATERIAL
    if modulo_elasticidade is None:
        modulo_elasticidade = 200000
    if coeficiente_poisson is None:
        coeficiente_poisson = 0.3
    if tensao_escoamento is None:
        tensao_escoamento = 534.1
    if tensao_ruptura is None:
        tensao_ruptura = 754.18
    
    tabela1 = doc.add_table(rows = 5, cols = 2)
    tabela1.style = estilo_tabela
    tabela1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela1.rows[0].cells[0].text = 'PROPRIEDADES DO MATERIAL'
##    tabela1.rows[1].cells[0].text = u'Módulo de Elasticidade (MPa)'
    lista_de_tab = [u'Módulo de Elasticidade (','E', ') (MPa)']
    paragrafo_celula = tabela1.rows[1].cells[0].add_paragraph()
    paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(len(lista_de_tab)):
        run = paragrafo_celula.add_run(lista_de_tab[i])
        if i in [1]:
            run.font.italic = True
##            if i == 2:
##                run.font.subscript = True
        run.font.name = 'Times New Roman'
##    tabela1.rows[2].cells[0].text = 'Coeficiente de Poisson'
    lista_de_tab = [u'Coeficiente de Poisson (','v', ')']
    paragrafo_celula = tabela1.rows[2].cells[0].add_paragraph()
    paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(len(lista_de_tab)):
        run = paragrafo_celula.add_run(lista_de_tab[i])
        if i in [1]:
            run.font.italic = True
##            if i == 2:
##                run.font.subscript = True
        run.font.name = 'Times New Roman'
##    tabela1.rows[3].cells[0].text = u'Tensão de Escoamento (MPa)'
    lista_de_tab = [u'Tensão de Escoamento (','S', 'yield', ') (MPa)']
    paragrafo_celula = tabela1.rows[3].cells[0].add_paragraph()
    paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(len(lista_de_tab)):
        run = paragrafo_celula.add_run(lista_de_tab[i])
        if i in [1, 2]:
            run.font.italic = True
            if i == 2:
                run.font.subscript = True
        run.font.name = 'Times New Roman'
##    tabela1.rows[4].cells[0].text = u'Tensão de Ruptura (MPa)'
    lista_de_tab = [u'Tensão de Ruptura (','S', 'ult', ') (MPa)']
    paragrafo_celula = tabela1.rows[4].cells[0].add_paragraph()
    paragrafo_celula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i in range(len(lista_de_tab)):
        run = paragrafo_celula.add_run(lista_de_tab[i])
        if i in [1, 2]:
            run.font.italic = True
            if i == 2:
                run.font.subscript = True
        run.font.name = 'Times New Roman'
    #tabela1.rows[0].cells[1].text = ''
    tabela1.rows[1].cells[1].text = str(modulo_elasticidade)#str(format(modulo_elasticidade, ".2e"))
    tabela1.rows[2].cells[1].text = str(coeficiente_poisson)#str(format(coeficiente_poisson, ".2f"))
    tabela1.rows[3].cells[1].text = str(tensao_escoamento)#str(format(tensao_escoamento, ".2e"))
    tabela1.rows[4].cells[1].text = str(tensao_ruptura)#str(format(tensao_ruptura, ".2e"))

    tabela1 = configuraTabela(doc, tabela1, 5, 2, titulo_unico = True)
    doc.add_paragraph()
    #GRÁFICO
    doc = configuraImagem(doc, arquivo_grafico, legenda = legenda1, flag_tam = 1)

    paragrafo2 = doc.add_paragraph(legenda3)
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo2.runs[0].font.name = 'Times New Roman'
    paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)

    #TABELA 2 - TENSÃO E DEFORMAÇÃO1
    if tensoes is None:
        tensoes = [0.0, 534.1, 586.356, 607.3325, 631.703, 642.6945, 653.7068, 670.2555, 675.3539, 680.4538, 685.5568, 690.6619, 700.875, 705.984, 718.6498, 721.3181, 729.7026, 740.6435, 754.1836]
    if deformacoes is None:
        deformacoes = [0.0, 0.002671, 0.009, 0.0125, 0.01872, 0.022567, 0.027259, 0.036259, 0.039593, 0.043232, 0.047203, 0.051533, 0.061388, 0.06698, 0.083037, 0.086861, 0.1, 0.12, 0.15]
    
    tabela2 = doc.add_table(rows = (len(tensoes) + 1), cols = 2)
    tabela2.style = estilo_tabela
    tabela2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela2.rows[0].cells[0].text = u'TENSÃO (MPa)'
    tabela2.rows[0].cells[1].text = u'DEFORMAÇÃO'
    for i in range(1, len(tensoes) + 1):
        tabela2.rows[i].cells[0].text = str(tensoes[i - 1])#str(format(tensoes[i - 1], ".2e"))
        tabela2.rows[i].cells[1].text = str(deformacoes[i - 1])#str(format(deformacoes[i - 1], ".2e"))

    tabela2 = configuraTabela(doc, tabela2, len(tensoes) + 1, 2, titulo_unico = False, negrito=False)#, legenda = legenda3) 
    doc.add_paragraph()
    return doc, num_figura, num_tabela

def modeloDoDefeitoIsolado(doc, num_topico, num_subtopico, num_tabela, num_figura,  profundidade = None,
                           comprimento = None, largura = None, posicao_supeprficie = None, pos_horaria = None,
                           pos_longitudinal = None, arquivo_imagem1 = None, arquivo_imagem2 = None, arquivo_imagem3 = None,
                           tr = None, fr = None):

    titulo = str(num_topico) + '.' + str(num_subtopico) + u' CARACTERÍSTICAS DO DEFEITO DE CORROSÃO ISOLADO'
    texto = [u'     As principais características geométricas dos defeitos (profundidade ', 'd', ', comprimento ', 'L', ' e largura ', 'W',
                 u', vide Figura ' + str(num_figura) + u') a serem analisados estão descritas na Tabela 5. Da mesma forma, estão indicados a posição do defeito na superfície do duto, ',
                 u'bem como os raios de concordância no topo do defeito (TR) e na frente do defeito (FR).']
##    texto = [u'    As principais características geométricas dos defeitos (profundidade ',
##                 'D', ', comprimento ', 'LL', ' e largura ', 'LC',
##                 u', vide figura) a serem analisadas estão descritas na tabela. Além do mais, está indicada a posição do defeito na superfície do duto.']

    legenda1 = 'Tabela ' + str(num_tabela) + u': Características geométricas e posições do defeito de corrosão'
    num_tabela += 1
    #legenda2 = 'Figura ' + str(num_figura) + ': Detalhe da discretização de um quarto do defeito.'
    legenda3 = 'Figura ' + str(num_figura) + u': Detalhe da geometria do defeito.'
    #legenda4 = 'Figura ' + str(num_figura + 1) + ': Vista da discretização do duto na região do defeito.'
##    num_figura += 1

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO
    paragrafo2 = doc.add_paragraph()
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i in range(len(texto)):
        run = paragrafo2.add_run(texto[i])
        if i in [1, 3, 5]:
            run.font.italic = True
            run.font.name = 'Times New Roman'

        run.font.size = Pt(tamanho_fonte_texto)

        #FIGURAS
    arquivo_imagem2 = os.path.join(os.getcwd(), 'GeradorRelatorio', 'esquematico_defeito_retangular.png')
    if arquivo_imagem1 is None:
        arquivo_imagem1 = 'defeito_1.jpg'
    if arquivo_imagem2 is None:
        arquivo_imagem2 = os.path.join(os.getcwd(), 'GeradorRelatorio', 'DimDefeitoRedimensionada.png')
    if arquivo_imagem3 is None:
        arquivo_imagem3 = 'defeito_isolado.png'
##    print('certo_final')
    #doc = configuraImagem(doc, arquivo_imagem1, legenda = legenda2)
    doc = configuraImagem(doc, arquivo_imagem2, legenda = legenda3)
    #doc = configuraImagem(doc, arquivo_imagem3, legenda = legenda4)

    #TABELA
    paragrafo2 = doc.add_paragraph(legenda1)
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo2.runs[0].font.name = 'Times New Roman'
    paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)

##    legenda = legenda1
    if profundidade is None:
        profundidade = 600
    if comprimento is None:
        comprimento = 10
    if largura is None:
        largura = 5
    if posicao_supeprficie is None:
        posicao_supeprficie = 'Externo'
    if pos_horaria is None:
        pos_horaria = '06:00:00'
    if pos_longitudinal is None:
        pos_longitudinal = '02:00:00'

    tabela = doc.add_table(rows = 10, cols = 2)
    tabela.style = estilo_tabela
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.rows[0].cells[0].text = u'DADOS DO DEFEITO DE CORROSÃO ISOLADO'
    tabela.rows[1].cells[0].text = 'Profundidade (mm)'
    tabela.rows[2].cells[0].text = 'Comprimento (mm)'
    tabela.rows[3].cells[0].text = 'Largura (mm)'
    tabela.rows[4].cells[0].text = u'Posição na Superfície do Duto'
    tabela.rows[5].cells[0].text = u'Posição Horária do Defeito no Duto (graus)(*)'
    tabela.rows[6].cells[0].text = u'Posição Longitudinal do Defeito no Duto (mm)(*)'
    tabela.rows[7].cells[0].text = u'Raio de concordância (Top fillet radius: TR)(mm)'
    tabela.rows[8].cells[0].text = u'Raio de concordância (Front fillet radius: FR)(mm)'
    #tabela.rows[7].cells[0].text = ''
    #tabela.rows[0].cells[1].text = ''
    tabela.rows[1].cells[1].text = str(profundidade)
    tabela.rows[2].cells[1].text = str(comprimento)
    tabela.rows[3].cells[1].text = str(largura)
    tabela.rows[4].cells[1].text = str(posicao_supeprficie)
    tabela.rows[5].cells[1].text = str(pos_horaria)
    tabela.rows[6].cells[1].text = str(pos_longitudinal)
    tabela.rows[7].cells[1].text = str(tr)
    tabela.rows[8].cells[1].text = str(fr)
    tabela.rows[9].cells[1].text = u'(*) Medida a partir do canto superior direito do defeito (vide Figura ' + str(num_figura) + ')'

    tabela = configuraTabela(doc, tabela, 10, 2, titulo_unico = True)

    cell90 = tabela.rows[9].cells[0]
    cell91 = tabela.rows[9].cells[1]
    cell90.merge(cell91)
    cell90.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    #para separar a tabela das imagens
    doc.add_paragraph()

    num_figura += 1
    

    return doc, num_figura, num_tabela

def modeloParaMultiplosDefeitos(doc, num_topico, num_subtopico, num_tabela, num_figura,
                                qtd_defeitos = None, profundidades = None, comprimentos = None,
                                larguras = None, posicoes_superficie = None, posicoes_horarias = None,
                                posicoes_longitudinais = None, arquivo_imagem1 = None, arquivo_imagem2 = None, arquivo_imagem3 = None):

    titulo = '\n' + str(num_topico) + '.' + str(num_subtopico) + ' MODELO PARA MÚLTIPLOS DEFEITOS'
    texto1 = [u'    As principais características geométricas dos defeitos (profundidade ', 'D',
                  ', comprimento ', 'LL', ' e largura ', 'LC',
                  u', vide figura) a serem analisadas estão descritas na tabela. Além do mais, está indicada a posição do defeito na superfície do duto.']
    texto2 = u'    Na situação de múltiplos defeitos,  a posição horária e posição longitudinal dos defeitos são informadas na tabela (ambas medidas a partir do canto superior esquerdo do defeito).'

    #legenda2 = 'Figura ' + str(num_figura) + ': Detalhe da discretização de um quarto do defeito.'
    legenda3 = 'Figura ' + str(num_figura) + u': Detalhe da geometria do defeito.'
    #legenda4 = 'Figura ' + str(num_figura + 1) + ': Vista da discretização do duto na região do defeito.'
    num_figura += 1

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO 1
    paragrafo2 = doc.add_paragraph()
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i in range(len(texto1)):
        run = paragrafo2.add_run(texto1[i])
        if i in [1, 3, 5]:
            run.font.italic = True
            
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    #TEXTO 2
    doc = configuraTexto(doc, texto2)

    #TABELAS
    if qtd_defeitos is None:
        qtd_defeitos = 2
    if profundidades is None:
        profundidades = [600, 300]
    if comprimentos is None:
        comprimentos = [10, 20]
    if larguras is None:
        larguras = [5, 10]
    if posicoes_superficie is None:
        posicoes_superficie = ['Externo', 'Interno']
    if posicoes_horarias is None:
        posicoes_horarias = [10, 15]
    if posicoes_longitudinais is None:
        posicoes_longitudinais = [10, 15]
    for i in range(qtd_defeitos):
        tabela = doc.add_table(rows = 8, cols = 2)
        tabela.style = estilo_tabela
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
        tabela.rows[0].cells[0].text = 'DADOS DO DEFEITO ' + str(i + 1)
        tabela.rows[1].cells[0].text = 'Profundidade (mm)'
        tabela.rows[2].cells[0].text = 'Comprimento (mm)'
        tabela.rows[3].cells[0].text = 'Largura (mm)'
        tabela.rows[4].cells[0].text = u'Posição na Superfície do Duto'
        tabela.rows[5].cells[0].text = u'Posição Horária do Defeito no Duto (graus)*'
        tabela.rows[6].cells[0].text = u'Posição Longitudinal do Defeito no Duto (mm)*'
        #tabela.rows[7].cells[0].text = ''        
        #tabela.rows[0].cells[1].text = ''
        tabela.rows[1].cells[1].text = str(format(profundidades[i], ".2e"))
        tabela.rows[2].cells[1].text = str(format(comprimentos[i], ".2e"))
        tabela.rows[3].cells[1].text = str(format(larguras[i], ".2e"))
        tabela.rows[4].cells[1].text = str(posicoes_superficie[i])
        tabela.rows[5].cells[1].text = str(posicoes_horarias[i])
        tabela.rows[6].cells[1].text = str(format(posicoes_longitudinais[i], ".2e"))
        tabela.rows[7].cells[1].text = '*Medida a partir do canto superior esquerdo do defeito.'

        legenda = 'Tabela ' + str(num_tabela) + u': Características geométricas e posições do defeito ' + str(i + 1) + '.'
        num_tabela += 1

        tabela = configuraTabela(doc, tabela, 8, 2, titulo_unico = True, altura_linha = 0.85, legenda = legenda)

        cell70 = tabela.rows[7].cells[0]
        cell71 = tabela.rows[7].cells[1]
        cell70.merge(cell71)
        cell70.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        doc.add_paragraph()        

    #FIGURAS
    if arquivo_imagem1 is None:
        arquivo_imagem1 = os.path.join(os.getcwd(), 'GeradorRelatorio', 'defeito_1.jpg')
    if arquivo_imagem2 is None:
        arquivo_imagem2 = os.path.join(os.getcwd(), 'GeradorRelatorio', 'DimDefeitoRedimensionada.png')
    if arquivo_imagem3 is None:
        arquivo_imagem3 = os.path.join(os.getcwd(), 'GeradorRelatorio', 'multiplos_defeitos.png')

    #doc = configuraImagem(doc, arquivo_imagem1, legenda = legenda2)
    doc = configuraImagem(doc, arquivo_imagem2, legenda = legenda3)
    #doc = configuraImagem(doc, arquivo_imagem3, legenda = legenda4)

    return doc, num_figura, num_tabela

def discretizacaoDoModeloDeDutoCorroido(doc, num_topico, num_subtopico, num_figura,  arquivo_discr_um_quarto = None, arquivo_discr_prox_defeito = None, model=None):

    titulo = str(num_topico) + '.' + str(num_subtopico) + u' DISCRETIZAÇÃO DO MODELO DE DUTO CORROÍDO'
    if model == '3D':
        texto1 = u'    De posse do modelo geométrico definido, a malha discreta é gerada fazendo-se uso de elementos sólidos (SOLID45) tridimensionais lineares ao longo de todo modelo.'

        ####VER ESCRITA DO TEXTO PARA O CASO DO DUTO COM DEFEITO REAL
        texto2 = u'    A Figura ' + str(num_figura) + u' mostra a malha de elementos finitos (EF) do modelo do segmento de duto com defeito idealizado e isolado (1/4 do modelo por conta da simetria) gerado pelo ANSYS.'
    else:
        texto1 = u'    De posse do modelo geométrico definido, a malha discreta é gerada fazendo-se uso de elementos planos (PLANE82) bidimensionais lineares ao longo de todo modelo.'

        ####VER ESCRITA DO TEXTO PARA O CASO DO DUTO COM DEFEITO REAL
        texto2 = u'    A Figura ' + str(num_figura) + u' mostra a malha de elementos finitos (EF) do modelo do segmento de duto com defeito isolado (perfil do modelo por conta da simetria) gerado pelo ANSYS.'
    
    legenda1 = 'Figura ' + str(num_figura) + u': Malha de EF do segmento de duto com defeito.'
    #legenda2 = 'Figura ' + str(num_figura + 1) + ': Vista da discretização do duto na região próxima do defeito.'
    num_figura += 1

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO 1
    doc = configuraTexto(doc, texto1)

    #TEXTO 2
    doc = configuraTexto(doc, texto2)

    #FIGURAS
    if arquivo_discr_um_quarto is None:
        arquivo_discr_um_quarto = os.path.join(os.getcwd(), 'GeradorRelatorio', 'discretizacao_um_quarto.png')
    if arquivo_discr_prox_defeito is None:
        arquivo_discr_prox_defeito = os.path.join(os.getcwd(), 'GeradorRelatorio', 'discretizacao_prox_defeito.png')

    doc = configuraImagem(doc, arquivo_discr_um_quarto, legenda = legenda1, flag_tam = 1)
    #doc = configuraImagem(doc, arquivo_discr_prox_defeito, legenda = legenda2)

    return doc, num_figura

def condicoesDeContorno(doc, num_topico, num_subtopico, num_figura,  tipo_de_duto = None, arquivo_restricao_z = None, arquivo_restricao_x = None, arquivo_restricao_y = None, arquivo_pressao_longitudinal = None, momento=None):

    if tipo_de_duto is None:
        tipo_de_duto = 'restrito'

    if tipo_de_duto == 'restrito':
        tipo_de_duto_txt = 'restrito axialmente'

    elif tipo_de_duto == 'nao restrito com tampa':
        tipo_de_duto_txt = u'não restrito com tampa'

    elif tipo_de_duto == 'nao restrito':
        tipo_de_duto_txt = u'não restrito'

    #para todos os tipos de dutos, esse texto se repete
    titulo = str(num_topico) + '.' + str(num_subtopico) + u' CONDIÇÕES DE CONTORNO'
    texto1 = u'    Para este modelo foi adotada a opção de duto ' + tipo_de_duto_txt + '.'
    texto2 = (u'    As condições de contorno aplicadas no modelo do segmento de duto com defeito, '
                  u'considerando-se as simetrias transversal e longitudinal, estão indicadas na Figura ' + str(num_figura) + '. '
                  u'As mesmas incluem: restrições nas direções longitudinal (Fix Z) e circunferencial (Fix X), '
                  u'respectivamente, nos planos de simetria transversal e longitudinal, mais um ponto da superfície '
                  u'interna, na extremidade longe, fixo na direção Y (Fix Y) para evitar movimento de corpo rígido '
                  u'do modelo nesta direção.')

    if float(momento) != float(0.0):
        texto2 += (u'Além disso, também são empregados na seção transversal da extremidade '
                   u'longe do defeito (região na cor verde) elementos finitos de casca (SHELL63) para fins de '
                   u'aplicação de carregamento de momento fletor.')

    #legendas que irão variar com o tipo de duto
    legenda = 'Figura ' + str(num_figura) + ': ' + u'Condições de contorno nos planos de simetria (transversal e longitudinal)'
    legenda1 = u'Restrição dos nós na direção Z (simetria e condições de extremidade).'
    legenda2 = u'Restrição dos nós na direção X (movimento de corpo livre e simetria).'
    legenda3 = u'Restrição dos nós na direção Y (evitar movimento de corpo livre).'
    legenda4 = u'Pressão longitudinal (na extremidade sem defeito).'

    #endereços das imagens que irão variar com o tipo de duto
    if arquivo_restricao_z is None:
        arquivo_restricao_z = os.path.join(os.getcwd(), 'GeradorRelatorio', 'duto_restrito_c_tampa_z.png')
    if arquivo_restricao_x is None:
        arquivo_restricao_x = os.path.join(os.getcwd(), 'GeradorRelatorio', 'duto_restrito_c_tampa_x.png')
    if arquivo_restricao_y is None:
        arquivo_restricao_y = os.path.join(os.getcwd(), 'GeradorRelatorio', 'duto_restrito_c_tampa_y.png')
    if arquivo_pressao_longitudinal is None:
        arquivo_pressao_longitudinal = os.path.join(os.environ['temp'], 'duto_nao_restrito_c_tampa.png')

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO 1
    doc = configuraTexto(doc, texto1)

    #TEXTO 2
    doc = configuraTexto(doc, texto2)

    arquivo_cond_cont = os.path.join(os.environ['temp'], "cond_contorno.png")
    if tipo_de_duto == 'restrito':
                
        legenda = 'Figura ' + str(num_figura) + ': ' + u'Condições de contorno, considerando 2 planos de simetria.'

        background = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "base.png"))
        img_z = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "z.png"))
        img_x = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "x.png"))
        img_y = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "y.png"))

        background.paste(img_z, (0, 0), img_z)
        background.paste(img_x, (0, 0), img_x)
        background.paste(img_y, (0, 0), img_y)
        background.save(arquivo_cond_cont, "PNG")
        doc = configuraImagem(doc, arquivo_cond_cont, legenda = legenda)
        num_figura += 1

    elif tipo_de_duto == 'nao restrito':
        
        legenda = 'Figura ' + str(num_figura) + ': ' + u'Condições de contorno, considerando 2 planos de simetria.'
        background = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "base.png"))
        img_z = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "z.png"))
        img_x = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "x.png"))
        img_y = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "y.png"))
        shell = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "shell.png"))

    if float(momento) != float(0.0):
        background.paste(shell, (0, 0), shell)
        background.paste(img_x, (0, 0), img_x)
        background.paste(img_y, (0, 0), img_y)
        background.paste(img_z, (0, 0), img_z)
        background.save(arquivo_cond_cont,"PNG")
        doc = configuraImagem(doc, arquivo_cond_cont, legenda = legenda)
        num_figura += 1
    
    elif tipo_de_duto == 'nao restrito com tampa':
        legenda = 'Figura ' + str(num_figura) + ': ' + u'Condições de contorno, considerando 2 planos de simetria.'
        background = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "base.png"))
        img_x = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "x.png"))
        img_y = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "y.png"))
        img_z = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "z.png"))
##        img_pressao = Image.open(os.path.join(os.getcwd(), 'GeradorRelatorio', "pressao_long.png"))
        background.paste(img_x, (0, 0), img_x)
        background.paste(img_y, (0, 0), img_y)
        background.paste(img_z, (0, 0), img_z)
##        background.paste(img_pressao, (0, 0), img_pressao)
        background.save(arquivo_cond_cont,"PNG")
        doc = configuraImagem(doc, arquivo_cond_cont, legenda = legenda)
        num_figura += 1

    return doc, num_figura

def cargasAplicadas(doc, num_topico, num_subtopico, num_figura, pressao_int, pressao_ext, temperatura, momento, forca_axial,  tipo_de_duto = None, arquivo_pressao_interna = None, arquivo_pressao_externa = None, arquivo_momento = None, arquivo_temperatura = None, arquivo_forca_axial = None):

    if tipo_de_duto is None:
        tipo_de_duto = 'restrito'

    titulo = str(num_topico) + '.' + str(num_subtopico) + ' CARGAS APLICADAS'
    texto = u'    Foram aplicadas as seguintes cargas: '
    texto1 = u'    Foi aplicada a seguinte carga: '
    texto2 = u'    Nenhuma carga foi aplicada no caso em análise.'

    texto3 = u'uma pressão interna de '
    texto4 = u'uma pressão externa de '
    texto5 = u'uma diferença de temperatura de '
    texto6 = u'um momento fletor de '
    texto7 = u'uma força axial de '

    texto8 = u' A figura a seguir indica o carregamento que foi aplicado no modelo.'
    texto9 = u' As figuras a seguir indicam os carregamentos que foram aplicados no modelo.'

    #arquivos das imagens
    if arquivo_pressao_interna is None:
        arquivo_pressao_interna = os.path.join(os.getcwd(), 'GeradorRelatorio', 'pressao_interna_valor_fornecido.png')
    if arquivo_pressao_externa is None:
        arquivo_pressao_externa = os.path.join(os.getcwd(), 'GeradorRelatorio', 'pressao_externa_valor_fornecido.png')
    if arquivo_momento is None:
        arquivo_momento = os.path.join(os.getcwd(), 'GeradorRelatorio', 'momento_valor_fornecido.png')
    if arquivo_temperatura is None:
        arquivo_temperatura = os.path.join(os.getcwd(), 'GeradorRelatorio', 'temperatura.png')
    if arquivo_forca_axial is None:
        arquivo_forca_axial = os.path.join(os.getcwd(), 'GeradorRelatorio', 'forca_axial.png')

    #legendas das imagens
    legenda1 = u'Pressão interna.'
    legenda2 = u'Pressão externa.'
    legenda3 = u'Momento fletor (ângulo do momento = 10°).'
    legenda4 = u'Definição do valor da pressão fornecido.'
    legenda5 = u'Temperatura (coeficiente de dilatação linear = 1,2 x 10e-5).'
    legenda6 = u'Força axial.'
    legenda = u'Carregamentos de serviço.'

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    background = Image.open(os.path.join(os.getcwd(), '', "base2.png"))
    img_p_interna = Image.open(os.path.join(os.getcwd(), '', "pressao_interna.png"))
    img_p_externa = Image.open(os.path.join(os.getcwd(), '', "pressao_externa.png"))
    img_f_axial = Image.open(os.path.join(os.getcwd(), '', "forca_axial.png"))
    img_momento = Image.open(os.path.join(os.getcwd(), '', "momento.png"))

    if tipo_de_duto == 'restrito':

        cargas = [pressao_int, pressao_ext, temperatura]

        qtd_cargas = np.count_nonzero(cargas)

        if qtd_cargas == 0: #não foram aplicadas cargas

            doc = configuraTexto(doc, texto2)

        elif qtd_cargas == 1: #apenas uma carga foi aplicada

            legenda_imagem = 'Figura ' + str(num_figura) + ': ' + legenda

            if pressao_int != 0:

                texto_final = texto1 + texto3 + str(format(pressao_int, ".2f")) + 'MPa.' + texto8

                doc = configuraTexto(doc, texto_final)

                background.paste(img_p_interna, (0, 0), img_p_interna)

            if pressao_ext != 0:

                texto_final = texto1 + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.' + texto8

                doc = configuraTexto(doc, texto_final)

                background.paste(img_p_externa, (0, 0), img_p_externa)

            if temperatura != 0:

                texto_final = texto1 + texto5 + str(format(temperatura, ".2f")) + '°C.'

                doc = configuraTexto(doc, texto_final)

            background.save(os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"),"PNG")
            doc = configuraImagem(doc, os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"), legenda = legenda_imagem)
            num_figura += 1

        else: #duas ou três cargas foram aplicadas
            contador = 0

            #MONTAR O TEXTO
            if pressao_int != 0:

                texto_final = texto + texto3 + str(format(pressao_int, ".2f")) + 'MPa'
                contador += 1

            if pressao_ext != 0:

                if contador == qtd_cargas - 1: #se for a última carga
                    texto_final = texto_final + ' e ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.'
                elif contador == 0: #se for a primeira carga
                    texto_final = texto + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'
                    contador += 1
                else: #se for a carga intermediária
                    texto_final = texto_final + ', ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'

            if temperatura != 0:

                texto_final = texto_final + ' e ' + texto5 + str(format(temperatura, ".2f")) + '°C.'

            #TEXTO
            texto_final = texto_final + texto9
            doc = configuraTexto(doc, texto_final)

            #IMAGENS
            legenda_imagem = 'Figura ' + str(num_figura) + ': ' + legenda

            if pressao_int != 0:

                background.paste(img_p_interna, (0, 0), img_p_interna)

            if pressao_ext != 0:

                background.paste(img_p_externa, (0, 0), img_p_externa)
            
            background.save(os.path.join(os.getcwd(), '', "cargas_aplicadas.png"),"PNG")
            doc = configuraImagem(doc, os.path.join(os.getcwd(), '', "cargas_aplicadas.png"), legenda = legenda_imagem)
            num_figura += 1

    elif tipo_de_duto == 'nao restrito':
        
        cargas = [pressao_int, pressao_ext, momento, forca_axial]

        qtd_cargas = np.count_nonzero(cargas)

        if qtd_cargas == 0: #não foram aplicadas cargas

            doc = configuraTexto(doc, texto2)

        elif qtd_cargas == 1: #apenas uma carga foi aplicada

            legenda_imagem = 'Figura ' + str(num_figura) + ': ' + legenda

            if pressao_int != 0:

                texto_final = texto1 + texto3 + str(format(pressao_int, ".2f")) + 'MPa.' + texto8

                doc = configuraTexto(doc, texto_final)

                background.paste(img_p_interna, (0, 0), img_p_interna)

            if pressao_ext != 0:

                texto_final = texto1 + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.' + texto8

                doc = configuraTexto(doc, texto_final)

                background.paste(img_p_externa, (0, 0), img_p_externa)

            if momento != 0:

                texto_final = texto1 + texto6 + str(format(momento, ".2f")) + 'kN.m.' + texto8

                doc = configuraTexto(doc, texto_final)

                background.paste(img_momento, (0, 0), img_momento)

            if forca_axial != 0:

                texto_final = texto1 + texto7 + str(format(forca_axial, ".2f")) + 'kN.' + texto8

                doc = configuraTexto(doc, texto_final)

                background.paste(img_f_axial, (0, 0), img_f_axial)

            background.save(os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"),"PNG")
            doc = configuraImagem(doc, os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"), legenda = legenda_imagem)
            num_figura += 1

        else: #duas, três ou quatro cargas foram aplicadas
            contador = 0

            legenda_imagem = 'Figura ' + str(num_figura) + ': ' + legenda

            #MONTAR O TEXTO
            if pressao_int != 0:

                texto_final = texto + texto3 + str(format(pressao_int, ".2f")) + 'MPa'
                contador += 1

            if pressao_ext != 0:

                if contador == qtd_cargas - 1: #se for a última carga
                    texto_final = texto_final + ' e ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.'
                elif contador == 0: #se for a primeira carga
                    texto_final = texto + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'
                    contador += 1
                else: #se for a carga intermediária
                    texto_final = texto_final + ', ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'
                    contador += 1

            if momento != 0:

                if contador == qtd_cargas - 1: #se for a última carga
                    texto_final = texto_final + ' e ' + texto6 + str(format(momento, ".2f")) + 'kN.m.'
                elif contador == 0: #se for a primeira carga
                    texto_final = texto + texto6 + str(format(momento, ".2f")) + 'kN.m'
                    contador += 1
                else: #se for a carga intermediária
                    texto_final = texto_final + ', ' + texto6 + str(format(momento, ".2f")) + 'kN.m'
                    contador += 1

            if forca_axial != 0:

                texto_final = texto_final + ' e ' + texto7 + str(format(forca_axial, ".2f")) + 'kN.'

            #TEXTO
            texto_final = texto_final + texto9
            doc = configuraTexto(doc, texto_final)

            #IMAGENS
            if pressao_int != 0:

                background.paste(img_p_interna, (0, 0), img_p_interna)

            if pressao_ext != 0:

                background.paste(img_p_externa, (0, 0), img_p_externa)

            if momento != 0:

                background.paste(img_momento, (0, 0), img_momento)

            if forca_axial != 0:

                background.paste(img_f_axial, (0, 0), img_f_axial)

            background.save(os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"),"PNG")
            doc = configuraImagem(doc, os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"), legenda = legenda_imagem)
            num_figura += 1

    elif tipo_de_duto == 'nao restrito com tampa':
        
        cargas = [pressao_int, pressao_ext, forca_axial]

        qtd_cargas = np.count_nonzero(cargas)

        if qtd_cargas == 0: #não foram aplicadas cargas

            doc = configuraTexto(doc, texto2)

        elif qtd_cargas == 1: #apenas uma carga foi aplicada

            legenda_imagem = 'Figura ' + str(num_figura) + u': ' + legenda

            if pressao_int != 0:

                texto_final = texto1 + texto3 + str(format(pressao_int, ".2f")) + 'MPa.' + texto8

                doc = configuraTexto(doc, texto_final)

                background.paste(img_p_interna, (0, 0), img_p_interna)

            if pressao_ext != 0:

                texto_final = texto1 + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.' + texto8

                doc = configuraTexto(doc, texto_final)

                background.paste(img_p_externa, (0, 0), img_p_externa)

            if forca_axial != 0:

                texto_final = texto1 + texto7 + str(format(forca_axial, ".2f")) + 'kN.' + texto8

                doc = configuraTexto(doc, texto_final)

                background.paste(img_f_axial, (0, 0), img_f_axial)

            background.save(os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"), "PNG")
            doc = configuraImagem(doc, os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"), legenda = legenda_imagem)
            num_figura += 1

        else: #duas ou três cargas foram aplicadas

            contador = 0

            legenda_imagem = 'Figura ' + str(num_figura) + u': ' + legenda

            #MONTAR O TEXTO
            if pressao_int != 0:

                texto_final = texto + texto3 + str(format(pressao_int, ".2f")) + 'MPa'
                contador += 1

            if pressao_ext != 0:

                if contador == qtd_cargas - 1: #se for a última carga
                    texto_final = texto_final + ' e ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa.'
                elif contador == 0: #se for a primeira carga
                    texto_final = texto + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'
                    contador += 1
                else: #se for a carga intermediária
                    texto_final = texto_final + ', ' + texto4 + str(format(pressao_ext, ".2f")) + 'MPa'

            if forca_axial != 0:

                texto_final = texto_final + ' e ' + texto7 + str(format(forca_axial, ".2f")) + 'kN.'

            #TEXTO
            texto_final = texto_final + texto9
            doc = configuraTexto(doc, texto_final)

            #IMAGENS
            if pressao_int != 0:

                background.paste(img_p_interna, (0, 0), img_p_interna)

            if pressao_ext != 0:

                background.paste(img_p_externa, (0, 0), img_p_externa)

            if forca_axial != 0:

                background.paste(img_f_axial, (0, 0), img_f_axial)

            background.save(os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"),"PNG")
            doc = configuraImagem(doc, os.path.join(os.getcwd(), 'GeradorRelatorio', "cargas_aplicadas.png"), legenda = legenda_imagem)
            num_figura += 1

    return doc, num_figura

def analisePressaoFornecida(doc, num_topico, num_subtopico,  pressao_fornecida = None, arquivo_pressao_fornecida = None):

    if pressao_fornecida is None:
        pressao_fornecida = 10
    if arquivo_pressao_fornecida is None:
        arquivo_pressao_fornecida = os.path.join(os.getcwd(), 'GeradorRelatorio', 'pressao_fornecida.png')

    titulo = str(num_topico) + '.' + str(num_subtopico) + u' CRITÉRIO DE PARADA'
##    texto = u'    A figura abaixo indica que o sistema tentará executar a análise até o valor da pressão fornecida de ' + str(format(pressao_fornecida, ".2f")) + 'MPa.'
##    texto1 = u'    O sistema tentará executar a análise até o valor da pressão fornecida de ' + str(format(pressao_fornecida, ".2f")) + 'MPa.'
    texto1 = (u'\tHá diferentes critérios de parada para a análise do segmento de duto submetido à pressão interna dominante. '
                  u'No CORDUT é possível selecionar dois critérios, seja até um valor de pressão fornecida ou até a ruptura do duto '
                  u'corroído. No presente caso, foi solicitado que análise fosse realizada até a pressão interna atingir o valor de '
                  + str(format(pressao_fornecida, ".2f"))+ ' Mpa.')
    #legenda = 'Figura ' + str(num_figura) + ': Definição do valor da pressão fornecido.'

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO 
    doc = configuraTexto(doc, texto1)

    #FIGURA
    #doc = configuraImagem(doc, arquivo_pressao_fornecida, legenda = legenda)

    return doc

def analisePressaoRuptura(doc, num_topico, num_subtopico,  arquivo_pressao_ruptura = None):

    if arquivo_pressao_ruptura is None:
        arquivo_pressao_ruptura = os.path.join(os.getcwd(), 'GeradorRelatorio', 'pressao_ruptura.png')

    titulo = str(num_topico) + '.' + str(num_subtopico) + u' CRITÉRIO DE PARADA'
##    texto = u'    A figura abaixo indica que o sistema tentará executar a análise até o valor da pressão de ruptura.'
##    texto1 = u'    O sistema tentará executar a análise até o valor da pressão de ruptura.'
    texto1 = (u'\tHá diferentes critérios de parada para a análise do segmento de duto submetido à pressão interna dominante. '
                  u'No CORDUT é possível selecionar dois critérios, seja até um valor de pressão fornecida ou até a ruptura do duto '
                  u'corroído.')
    #legenda = 'Figura ' + str(num_figura) + ': Escolha para análise de ruptura.'

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO 
    doc = configuraTexto(doc, texto1)

    #FIGURA
    #doc = configuraImagem(doc, arquivo_pressao_ruptura, legenda = legenda)

    return doc

def resultadosEscoamento(doc, num_topico, num_subtopico, num_subsubtopico, num_figura,  arquivo_mapa_tensoes = None, arquivo_janela_CORDUT = None):

    titulo = str(num_topico) + '.' + str(num_subtopico) + '.' +  str(num_subsubtopico) + u' Tensões no início do escoamento'
    texto = [u'    Na figura ' + str(num_figura) + u', está apresentado o mapa de tensões equivalentes de ', u'von Mises ', u'(SEQV) do modelo, no passo de carga em que se iniciou o escoamento.']

    #arquivos das imagens
    if arquivo_mapa_tensoes is None:
        arquivo_mapa_tensoes = os.path.join(os.getcwd(), 'GeradorRelatorio', 'mapa_tensoes.png')
    if arquivo_janela_CORDUT is None:
        arquivo_janela_CORDUT = os.path.join(os.getcwd(), 'GeradorRelatorio', 'janela_CORDUT.png')

    #legendas das imagens
    legenda1 = 'Figura ' + str(num_figura) + u': Distribuição das tensões de von Mises (SEQV) no passo de carga em que se iniciou o escoamento.'
    #legenda2 = 'Figura ' + str(num_figura + 1) + ': Janela de resultado do CORDUT no início do escoamento.'

    num_figura += 1

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO 1
##    doc = configuraTexto(doc, texto)

    paragrafo2 = doc.add_paragraph()
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i in range(len(texto)):
        run = paragrafo2.add_run(texto[i])
        if i in [1]:
            run.font.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)


    #FIGURA 1
    doc = configuraImagem(doc, arquivo_mapa_tensoes, legenda = legenda1, flag_tam = 1)

    #FIGURA 2
    #doc = configuraImagem(doc, arquivo_janela_CORDUT, legenda = legenda2)    

    return doc, num_figura

def resultadosPressaoFornecida(doc, num_topico, num_subtopico, num_subsubtopico, num_figura, tens_rup, tensao_parada,  arquivo_tensao_pressao = None, arquivo_janela_resultado = None):

    titulo = str(num_topico) + '.' + str(num_subtopico) + '.' +  str(num_subsubtopico)  + u' Tensões na pressão interna de parada fornecida'
    texto = [u'    Na Figura ' + str(num_figura) + u', está apresentada a perspectiva do mapa de tensões equivalentes de ', 'von Mises ',
                 u'do modelo do duto com defeito de corrosão, no passo de carga que atingiu a pressão de parada fornecida. A tensão máxima de ', 'von Mises ',
                 u'é de ' + str(tensao_parada) + u' MPa ('+ str(format(100*tensao_parada/tens_rup, '.2f')) + u'% da tensão última verdadeira).']

    #legendas das imagens
    legenda1 = 'Figura ' + str(num_figura) + u': Distribuição das tensões de von Mises (SEQV) no passo de carga que atingiu a pressão de parada fornecida'
    #legenda2 = 'Figura ' + str(num_figura + 1) + ': Janela de resultado do CORDUT no passo de carga que atingiu a pressão fornecida.'
    num_figura += 1

    #arquivo das imagens
    if arquivo_tensao_pressao is None:
        arquivo_tensao_pressao = os.path.join(os.getcwd(), 'GeradorRelatorio', 'tensao_pressao.png')
    if arquivo_janela_resultado is None:
        arquivo_janela_resultado = os.path.join(os.getcwd(), 'GeradorRelatorio', 'janela_CORDUT_resultado.png')

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

##    #TEXTO 1
##    doc = configuraTexto(doc, texto)
    paragrafo2 = doc.add_paragraph()
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i in range(len(texto)):
        run = paragrafo2.add_run(texto[i])
        if i in [1, 3]:
            run.font.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    #FIGURA 1
    doc = configuraImagem(doc, arquivo_tensao_pressao, legenda = legenda1, flag_tam = 1)

    #FIGURA 2
    #doc = configuraImagem(doc, arquivo_janela_resultado, legenda = legenda2)

    return doc, num_figura

def resultadosPressaoRuptura(doc, num_topico, num_subtopico, num_subsubtopico, num_figura,  arquivo_tensao_pressao = None, arquivo_janela_resultado = None):

    titulo = str(num_topico) + '.' + str(num_subtopico) + '.' +  str(num_subsubtopico)  + u' Tensões na ruptura'
    texto = u'    Na Figura ' + str(num_figura) + u', está apresentada a perspectiva do mapa de tensões do modelo do duto com defeito na ruptura.'
    texto = [u'    Na Figura ' + str(num_figura) + u', está apresentada a perspectiva do mapa de tensões equivalentes de ', 'von Mises ',
                 u'do modelo do duto com defeito, no passo de carga que atingiu a pressão de parada.']

    #legenda das imagens
    legenda1 = 'Figura ' + str(num_figura) + u': Distribuição das tensões de von Mises (SEQV) no passo de carga que atingiu a pressão de parada'
    #legenda2 = 'Figura ' + str(num_figura + 1) + ': Janela de resultado do CORDUT no final da análise.'
    num_figura += 1

    #arquivo das imagens
    if arquivo_tensao_pressao is None:
        arquivo_tensao_pressao = os.path.join(os.getcwd(), 'GeradorRelatorio', 'tensao_pressao_ruptura.png')
    if arquivo_janela_resultado is None:
        arquivo_janela_resultado = os.path.join(os.getcwd(), 'GeradorRelatorio', 'janela_CORDUT_resultado_ruptura.png')

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO 1
##    doc = configuraTexto(doc, texto)
    paragrafo2 = doc.add_paragraph()
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i in range(len(texto)):
        run = paragrafo2.add_run(texto[i])
        if i in [1]:
            run.font.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    #FIGURA 1
    doc = configuraImagem(doc, arquivo_tensao_pressao, legenda = legenda1, flag_tam = 1)

    #FIGURA 2
    #doc = configuraImagem(doc, arquivo_janela_resultado, legenda = legenda2)

    return doc, num_figura

def deslocamentos(doc, num_topico, num_subtopico, num_subsubtopico, num_figura,  analise_ruptura = None, arquivo_deslocamento = None, arquivo_deslocamento_1 = None):

    if analise_ruptura is None:
        analise_ruptura = False
        titulo_text = u' Deslocamentos na pressão de parada fornecida'
    else:
                titulo_text = u' Deslocamentos na pressão de ruptura'

    titulo = str(num_topico) + '.' + str(num_subtopico) + '.' +  str(num_subsubtopico)  + titulo_text
    texto1 = u'    A deformada do duto na região do defeito (em escala automática do ANSYS)'
    texto2 = u', no passo de carga referente à pressão de parada fornecida'
    texto3 = u', no passo de carga referente à pressão de ruptura'
    texto4 = u', está ilustrada na Figura ' + str(num_figura) + u' e a deformada do duto na região do defeito (em escala real) está ilustrada na Figura ' + str(num_figura + 1) + '.'
    legenda1 = 'Figura ' + str(num_figura) + u': Configuração deformada do modelo do duto com defeito referente à pressão última (escala automática do ANSYS)'
    legenda3 = 'Figura ' + str(num_figura + 1) + u': Configuração deformada do modelo do duto com defeito referente à pressão última (escala real)'
    num_figura += 2
    legenda2 = '.'
    if analise_ruptura:
        texto = texto1 + texto3 + texto4
        legenda = legenda1 + texto3 + legenda2
        legenda_2 = legenda3 + texto3 + legenda2
    else:
        texto = texto1 + texto2 + texto4
        legenda = legenda1 + texto2 + legenda2
        legenda_2 = legenda3 + texto2 + legenda2

    #arquivo da imagem
    if arquivo_deslocamento is None:
        arquivo_deslocamento = os.path.join(os.getcwd(), 'GeradorRelatorio', 'deslocamento.png')

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO
    doc = configuraTexto(doc, texto)

    #FIGURA
    doc = configuraImagem(doc, arquivo_deslocamento, legenda = legenda, flag_tam = 1)

    doc = configuraImagem(doc, arquivo_deslocamento_1, legenda = legenda_2, flag_tam = 1)

    

    return doc, num_figura

def pressao_internar_deslocamento(doc, num_topico, num_subtopico, num_subsubtopico, num_figura, num_tabela, arquivo_pressao_deslocamento = None, press_desl = None):

    titulo = str(num_topico) + '.' + str(num_subtopico) + '.' +  str(num_subsubtopico)  + u' Gráfico e tabela relacionando pressão interna x deslocamento'
    texto = [u'    O histórico da aplicação incremental do carregamento de pressão interna ', 'versus ',
                  u'o deslocamento máximo (nó com o maior deslocamento no último passo) está apresentado na Figura',
                  u' ' + str(num_figura) + u', com respectivos valores contidos na Tabela 6.']

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO
    paragrafo2 = doc.add_paragraph()
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i in range(len(texto)):
        run = paragrafo2.add_run(texto[i])
        if i == 1:
            run.font.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    

    #GRÁFICO
    legenda1 = 'Figura '+ str(num_figura) + u': Curva pressão x deslocamento para o nó de deslocamento máximo'
    arquivo_grafico = arquivo_pressao_deslocamento
    doc = configuraImagem(doc, arquivo_grafico, legenda = legenda1, flag_tam = 1)

##    paragrafo2 = doc.add_paragraph(legenda3)
##    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
##    paragrafo2.runs[0].font.name = 'Times New Roman'
##    paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)

    paragrafo2 = doc.add_paragraph('Tabela ' + str(num_tabela) + u': Valores da Pressão interna x Deslocamento para o nó de deslocamento máximo')
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo2.runs[0].font.name = 'Times New Roman'
    paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)
    
##    #TABELA 1 - PROPRIEDADES DO MATERIAL
##    if modulo_elasticidade is None:
##        modulo_elasticidade = 200000
##    if coeficiente_poisson is None:
##        coeficiente_poisson = 0.3
##    if tensao_escoamento is None:
##        tensao_escoamento = 534.1
##    if tensao_ruptura is None:
##        tensao_ruptura = 754.18
    tabela1 = doc.add_table(rows = len(press_desl) + 1, cols = 2)
    tabela1.style = estilo_tabela
    tabela1.alignment = WD_TABLE_ALIGNMENT.CENTER

    tabela1.rows[0].cells[0].text = u'Pressão Interna'
    tabela1.rows[0].cells[1].text = u'Deslocamento'
    count = 1
    
    for (disp, press) in press_desl:
                tabela1.rows[count].cells[0].text = str(press)
                tabela1.rows[count].cells[1].text = str(disp)
                count += 1
##    tabela1.rows[0].cells[0].text = 'PROPRIEDADES DO MATERIAL'
##    tabela1.rows[1].cells[0].text = u'Módulo de Elasticidade (MPa)'
##    tabela1.rows[2].cells[0].text = 'Coeficiente de Poisson'
##    tabela1.rows[3].cells[0].text = u'Tensão de Escoamento (MPa)'
##    tabela1.rows[4].cells[0].text = u'Tensão de Ruptura (MPa)'
##    #tabela1.rows[0].cells[1].text = ''
##    tabela1.rows[1].cells[1].text = str(modulo_elasticidade)#str(format(modulo_elasticidade, ".2e"))
##    tabela1.rows[2].cells[1].text = str(coeficiente_poisson)#str(format(coeficiente_poisson, ".2f"))
##    tabela1.rows[3].cells[1].text = str(tensao_escoamento)#str(format(tensao_escoamento, ".2e"))
##    tabela1.rows[4].cells[1].text = str(tensao_ruptura)#str(format(tensao_ruptura, ".2e"))
##
    tabela1 = configuraTabela(doc, tabela1, len(press_desl)+1, 2, titulo_unico = False, negrito=False)
##
    doc.add_paragraph()
        
    num_tabela += 1
    num_figura += 1
    return (doc, num_figura, num_tabela)

def sh_sl(doc, num_topico, num_subtopico, num_subsubtopico, num_figura,  arquivo_sh_sl = None):

    titulo = str(num_topico) + '.' + str(num_subtopico) + '.' +  str(num_subsubtopico)  + u' Gráfico tensão circunferencial x tensão longitudinal'
    texto = [u'    A evolução da tensão circunferencial ', 'versus ',
                  u'tensão longitudinal, no nó com o maior deslocamento no último passo (curva verde) e no nó com o menor deslocamento no último passo (curva azul)',
                  u' são observados na Figura ' + str(num_figura) + '.']

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO
    paragrafo2 = doc.add_paragraph()
    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for i in range(len(texto)):
        run = paragrafo2.add_run(texto[i])
        if i == 1:
            run.font.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(tamanho_fonte_texto)

    

    #GRÁFICO
    legenda1 = 'Figura '+ str(num_figura) + u': Curvas tensão circunferencial x tensão longitudinal, em um nó dentro do defeito (curva verde) e em um nó longe do defeito (curva azul)'
##    arquivo_grafico = arquivo_pressao_deslocamento
    doc = configuraImagem(doc, arquivo_sh_sl, legenda = legenda1, flag_tam = 1)

##    paragrafo2 = doc.add_paragraph(legenda3)
##    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
##    paragrafo2.runs[0].font.name = 'Times New Roman'
##    paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)

##    paragrafo2 = doc.add_paragraph(u'Valores da Pressão interna x Deslocamento para o nó de deslocamento máximo')
##    paragrafo2.alignment = WD_ALIGN_PARAGRAPH.CENTER
##    paragrafo2.runs[0].font.name = 'Times New Roman'
##    paragrafo2.runs[0].font.size = Pt(tamanho_fonte_legenda)

    return(doc, num_figura)
                  

def recomendacoes(doc, num_topico,  recomendacao_usuario = None):

    titulo = str(num_topico) + u'. RECOMENDAÇÕES'

    default1 = [u'    A fim de avaliar o comprometimento da integridade estrutural dos casos dos dutos analisados, os resultados da pressão de um duto íntegro (especificar valor) e com defeito indicam que o duto poderá ser mantido em operação com carga inferior à de ruptura calculada (especificar valor).',
                    u'    Entretanto, recomenda-se definir a segurança da estrutura através do emprego de um fator de segurança fornecido pelas normas de projeto ou até mesmo através da análise de confiabilidade. Esta permite que inspeções e reparos sejam programados de forma a garantir a segurança do duto corroído.']
    default2 = [u'    O CORDUT concluiu a simulação com sucesso.', u'    A tensão de escoamento corresponde a um valor de 534 Mpa.',
                    u'    A tensão referente a um valor de pressão fornecida corresponde a um valor de 620 Mpa.',
                    u'    Neste caso, a ruptura não foi detectada (tensão de ruptura 754,18 Mpa).', u'    O duto poderá manter em operação com segurança.']

    #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    #TEXTO
    '''
    if recomendacao_usuario is None:
        for frase in default1:
            doc = configuraTexto(doc, frase)
    else:
        doc = configuraTexto(doc, recomendacao_usuario)
    '''
    doc = configuraTexto(doc, '    ')

    return doc

def anexo(doc, de, t, ldefeito, w, d, lduto, superficie):
    titulo = u'ANEXO'
        #TÍTULO
    doc = configuraTexto(doc, titulo, titulo = True)

    if superficie == 'externa':
        w_grade = format((360*w/(math.pi * de)), ".2f")
    else:
        w_grade = format((360*w/(math.pi * (de-(2*t)))), ".2f")

    tabela = doc.add_table(rows = 10, cols = 2)
    tabela.style = estilo_tabela
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabela.rows[0].cells[0].text = ''
    tabela.rows[1].cells[0].text = 'De (mm)'
    tabela.rows[2].cells[0].text = 'De (pol)'
    tabela.rows[3].cells[0].text = 't (mm)'
    tabela.rows[4].cells[0].text = u'd/t'
    tabela.rows[5].cells[0].text = u'L (mm)'
    tabela.rows[6].cells[0].text = u'W (mm)'
    tabela.rows[7].cells[0].text = u'W (graus)'
    tabela.rows[8].cells[0].text = u'd (mm)'
    tabela.rows[9].cells[0].text = u'Lduto (mm)'
    #tabela.rows[7].cells[0].text = ''
    tabela.rows[0].cells[1].text = 'Ponto 1'
    tabela.rows[1].cells[1].text = str(de)
    tabela.rows[2].cells[1].text = str(format((de/25.4), ".2f"))
    tabela.rows[3].cells[1].text = str(t)
    tabela.rows[4].cells[1].text = str(format((d/t), ".2f"))
    tabela.rows[5].cells[1].text = str(ldefeito)
    tabela.rows[6].cells[1].text = str(w)
    tabela.rows[7].cells[1].text = str(w_grade)
    tabela.rows[8].cells[1].text = str(d)
    tabela.rows[9].cells[1].text = str(lduto)
    tabela = configuraTabela(doc, tabela, 10, 2, titulo_unico = False)

    return doc
        

def gerarRelatorio(dicionario, dir_out):

    #organizadores gerais do relatório
    num_topico = 1
    num_tabela = 1
    num_figura = 1

    #criação do documento
    doc = docx.Document()

        #inicialização dos dados de identificação do relatório
    doc = identificacaoDoRelatorio(doc, num_topico)

    #inicialização da quantidade de casos analisados no relatório
    num_topico += 1
    qtd_casos = len(dicionario)
    doc, num_tabela = quantidadeDeCasos(doc, num_topico, num_tabela, qtd_casos = qtd_casos)

    for i in range(qtd_casos):

        #obtendo as informações do caso a ser analisado
        caso = dicionario[i + 1]
        
        #obtendo as informações das imagens que irão compor o relatório
        imagens = caso['imagens']

        #obtendo as informações do duto
        informacoes_duto = caso['duto']

        #organizador de cada caso analisado do relatório
        num_subtopico = 1
        num_topico += 1
        doc, num_tabela = modeloDoDuto(doc, num_topico, num_subtopico, num_tabela, diametro_externo = informacoes_duto['diam_ext'], espessura = informacoes_duto['espessura'], comprimento = informacoes_duto['comp_duto'], taxa_ovalizacao = informacoes_duto['ovalizacao'], nome_modelo = caso['nome_modelo'])

        #obtendo as informações do material
        informacoes_material = caso['material']

        #obtendo as informações da curva tensão x deformação do material
        duplas_tens_def = informacoes_material['lista_curva_material']
        lista_tensoes = []
        lista_deformacoes = []
        for dupla in duplas_tens_def:
            lista_tensoes.append(dupla[0])
            lista_deformacoes.append(dupla[1])

        #propriedades do material
        num_subtopico += 1
        doc, num_figura, num_tabela = propriedadesDoMaterial(doc, num_topico, num_subtopico, num_tabela, num_figura, arquivo_grafico = informacoes_material['dir_arq_cf_plot'], modulo_elasticidade = informacoes_material['mod_elast'], coeficiente_poisson = informacoes_material['coef_poisson'], tensao_escoamento = informacoes_material['tens_esc'], tensao_ruptura = informacoes_material['tens_rup'], tensoes = lista_tensoes, deformacoes = lista_deformacoes)

        #obtendo as informações sobre os defeitos
        informacoes_defeitos = caso['defeitos']
        qtd_defeitos = len(informacoes_defeitos)

        if qtd_defeitos == 1: #DEFEITO ISOLADO

            #obtendo as informações sobre o defeito isolado
            defeito = informacoes_defeitos[1]

            num_subtopico += 1
            # if defeito['type'] == 'Idealized':
            #     doc, num_figura, num_tabela = modeloDoDefeitoIsolado(doc, num_topico, num_subtopico, num_tabela,
            #                                                          num_figura, profundidade = defeito['prof'],
            #                                                          comprimento = defeito['comp_def'],
            #                                                          largura = defeito['larg_def'], posicao_supeprficie = defeito['superficie'].capitalize(),
            #                                                          pos_horaria = defeito['pos_hor'], pos_longitudinal = defeito['pos_long'],
            #                                                          tr = defeito['TR'], fr = defeito['FR'])

        else: #MÚLTIPLOS DEFEITOS

            #organizando as informações dos defeitos
            lista_profundidades = []
            lista_comprimentos = []
            lista_larguras = []
            lista_pos_superficie = []
            lista_pos_hor = []
            lista_pos_long = []
            for i in range(qtd_defeitos):
                defeito = informacoes_defeitos[i + 1]

                lista_profundidades.append(defeito['prof'])
                lista_comprimentos.append(defeito['comp_def'])
                lista_larguras.append(defeito['larg_def'])
                lista_pos_superficie.append(defeito['superficie'].capitalize())
                lista_pos_hor.append(defeito['pos_hor'])
                lista_pos_long.append(defeito['pos_long'])

            num_subtopico += 1
            doc, num_figura, num_tabela = modeloParaMultiplosDefeitos(doc, num_topico, num_subtopico, num_tabela, num_figura, qtd_defeitos = qtd_defeitos, profundidades = lista_profundidades, comprimentos = lista_comprimentos, larguras = lista_larguras, posicoes_superficie = lista_pos_superficie, posicoes_horarias = lista_pos_hor, posicoes_longitudinais = lista_pos_long)
            
        # num_subtopico += 1
        # doc, num_figura = discretizacaoDoModeloDeDutoCorroido(doc, num_topico, num_subtopico, num_figura, arquivo_discr_um_quarto = imagens['disc_um_quarto_duto'], model=informacoes_duto['model'])
        
        # num_subtopico += 1
        # doc, num_figura = condicoesDeContorno(doc, num_topico, num_subtopico, num_figura, tipo_de_duto = caso['cond_ext'], momento = caso['carreg_serv']['momento'])

        #obtendo as informações sobre as cargas aplicadas
        cargas_aplicadas = caso['carreg_serv']
        num_subtopico += 1
        doc, num_figura = cargasAplicadas(doc, num_topico, num_subtopico, num_figura, pressao_int = cargas_aplicadas['pint'],
                                                  pressao_ext = cargas_aplicadas['pext'], temperatura = cargas_aplicadas['temp'],
                                                  momento = cargas_aplicadas['momento'], forca_axial = cargas_aplicadas['faxial'], tipo_de_duto = caso['cond_ext'])

        #obtendo as informações sobre a análise
        analise = caso['analise']

        tipo_analise = analise['condicao_parada']
        num_subsubtopico = 1
        if tipo_analise == 'press_fornecida':

            num_subtopico += 1
            doc = analisePressaoFornecida(doc, num_topico, num_subtopico, pressao_fornecida = analise['pressao_solicitante'])
            num_subtopico += 1
            doc = configuraTexto(doc, str(num_topico) + '.' + str(num_subtopico) + '. RESULTADOS', titulo = True)
            #terá escoamento se a pressão fornecida for maior ou igual a pressão de escoamento
            print('pressao_solicitante: ', analise['pressao_solicitante'])
            print('press_esc: ', analise['press_esc'])
            if float(analise['pressao_solicitante']) >= float(analise['press_esc']):
                print('pressao_solicitante: ', analise['pressao_solicitante'])
                print('press_esc: ', analise['press_esc'])
##                num_subtopico += 1
##                num_subsubtopico += 1
                doc, num_figura = resultadosEscoamento(doc, num_topico, num_subtopico, num_subsubtopico,
                                                                       num_figura, arquivo_mapa_tensoes = imagens['mapa_tensoes_inicio_escoamento'])
                
##            num_subtopico += 1
            num_subsubtopico += 1
            doc, num_figura = resultadosPressaoFornecida(doc, num_topico, num_subtopico, num_subsubtopico,
                                                                     num_figura, informacoes_material['tens_rup'], analise['tensao_parada'], arquivo_tensao_pressao = imagens['res_analise'])
            
##            num_subtopico += 1
            # num_subsubtopico += 1
            # doc, num_figura = deslocamentos(doc, num_topico, num_subtopico, num_subsubtopico,
            #                                 num_figura, analise_ruptura = False, arquivo_deslocamento = imagens['deslocamentos'],
            #                                 arquivo_deslocamento_1 = imagens['deslocamentos_1'])

        else:
                        #TÍTULO
                        
            num_subtopico += 1
            doc = analisePressaoRuptura(doc, num_topico, num_subtopico, num_subsubtopico)
            num_subtopico += 1
            doc = configuraTexto(doc, str(num_topico) + '.' + str(num_subtopico) + '. RESULTADOS', titulo = True)
            #quando há análise até ruptura, há resultado à nível de escoamento
##            num_subtopico += 1
##            num_subsubtopico += 1
            doc, num_figura = resultadosEscoamento(doc, num_topico, num_subtopico, num_subsubtopico, num_figura,
                                                               arquivo_mapa_tensoes = imagens['mapa_tensoes_inicio_escoamento'])
            
##            num_subtopico += 1
            num_subsubtopico += 1
            doc, num_figura = resultadosPressaoRuptura(doc, num_topico, num_subtopico, num_subsubtopico,
                                                                   num_figura, arquivo_tensao_pressao = imagens['res_analise'])
            
##            num_subtopico += 1
            num_subsubtopico += 1
            doc, num_figura = deslocamentos(doc, num_topico, num_subtopico, num_subsubtopico,
                                                        num_figura, analise_ruptura = True, arquivo_deslocamento = imagens['deslocamentos'],
                                                        arquivo_deslocamento_1 = imagens['deslocamentos_1'])

        # num_subsubtopico += 1
        # doc, num_figura, num_tabela = pressao_internar_deslocamento(doc, num_topico, num_subtopico, num_subsubtopico,
        #                                                             num_figura, num_tabela, arquivo_pressao_deslocamento = imagens['disp_press'],
        #                                                             press_desl = analise['press_desl'])
        # num_subsubtopico += 1
        # doc, num_figura = sh_sl(doc, num_topico, num_subtopico, num_subsubtopico, num_figura,  arquivo_sh_sl = imagens['sh_sl'])

    num_topico += 1
    doc = recomendacoes(doc, num_topico)

    # if defeito['type'] == 'Idealized':
    #     anexo(doc, informacoes_duto['diam_ext'], informacoes_duto['espessura'],
    #     defeito['comp_def'], defeito['larg_def'], defeito['prof'], informacoes_duto['comp_duto'],
    #     defeito['superficie'].lower())
#informacoes_duto['diam_ext'], espessura = informacoes_duto['espessura'], comprimento = informacoes_duto['comp_duto']
##    defeito['prof'],
##                                                                                     comprimento = defeito['comp_def'],
##                                                                                     largura = defeito['larg_def']
    #salvando o relatório
    doc.save('relatorio.docx')
    # convert('relatorio.docx')
    # doc.save(dir_out + os.sep + 'relatorio.docx')
    #abrindo o arquivo automaticamente
    #os.system('start relatorio.docx')

    return doc

##########           TESTE            ############
if __name__ == '__main__':
        dicionario = {1: {'nome_modelo': 'idts_013', 'duto': {'diam_ext': 458.8, 'espessura': 8.2,
                                                              'comp_duto': 2340, 'ovalizacao': 2},
                          # 'material': {'dir_arq_cf_plot': 'tensao_versus_deformacao.png',
                           'material': {'dir_arq_cf_plot': 'z.png',
                                       'mod_elast': 199962.5608386372, 'coef_poisson': 0.3,
                                       'tens_esc': 534.1, 'tens_rup':754.1836,
                                       'lista_curva_material': [(0.0, 0.0),(534.1, 0.002671),
                                                                (586.356, 0.009), (607.3325, 0.0125),
                                                                (631.703, 0.0187197), (642.6945, 0.0225671),
                                                                (653.7068, 0.0272592), (670.2555, 0.0362591),
                                                                (675.3539, 0.039593), (680.4538, 0.0432324),
                                                                (685.5568, 0.0472031), (690.6619, 0.0515325),
                                                                (700.875, 0.061388), (705.984, 0.0669801),
                                                                (718.6498, 0.0830373), (721.3181, 0.0868612),
                                                                (729.7026, 0.1), (740.6435, 0.12),
                                                                (754.1836, 0.15)]}, 'defeitos':
                          {1: {'tipo_def': 'retangular', 'comp_def': 50, 'larg_def': 50, 'prof': 4,
                               'superficie': 'interna', 'pos_long': 1145, 'pos_hor': '06:00:00'}},
                          'cond_ext': 'restrito', 'carreg_serv': {'pint': 10.0, 'pext': 2.0,
                                                                  'temp': 0.0, 'alfa': 0.0,
                                                                  'momento': 0.0, 'faxial': 0.0,
                                                                  'ang_momento': 0.0},
                          'analise': {'press_parada': 15, 'tensao_parada': 600, 'press_esc': 10,
                                      'tensao_esc': 534.1, 'condicao_parada': 'press_fornecida',
                                      'pressao_solicitante': 15.0},
                          'imagens': {'disc_um_quarto_duto': 'discretizacao_um_quarto.png',
                                      'mapa_tensoes_inicio_escoamento': 'mapa_tensoes.png',
                                      'res_analise': 'tensao_pressao.png', 'deslocamentos': 'deslocamento.png'}}}

        gerarRelatorio(dicionario, '')
