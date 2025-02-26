# -*- coding: utf-8 -*-
"""
Created on Wed Nov 20 18:17:29 2024

@author: MottaRS
"""


from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import pandas as pd

# Global parameters
FONT_NAME = 'Times New Roman'
FONT_SIZE_TITLE = 15
FONT_SIZE_TEXT = 12
FONT_SIZE_TABLE = 10
FONT_SIZE_CAPTION = 10
TABLE_STYLE = 'Light Grid Accent 1'

def configure_text(doc, text, is_title=False, level=0):
    """
    Add formatted text to the document.
    """
    paragraph = doc.add_paragraph(text)
    if is_title:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.runs[0].font.name = FONT_NAME
        paragraph.runs[0].font.size = Pt(FONT_SIZE_TITLE)
        paragraph.runs[0].font.bold = True
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.runs[0].font.name = FONT_NAME
        paragraph.runs[0].font.size = Pt(FONT_SIZE_TEXT)
    return doc

def create_table(doc, data):
    """
    Add a table to the document from a pandas DataFrame.
    """
    table = doc.add_table(rows=1, cols=len(data.columns), style=TABLE_STYLE)
    header_cells = table.rows[0].cells
    for i, column_name in enumerate(data.columns):
        header_cells[i].text = column_name

    for _, row in data.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    return doc

def add_image(doc, image_path, width=Cm(16)):
    """
    Add an image to the document.
    """
    try:
        doc.add_picture(image_path, width=width)
    except FileNotFoundError:
        doc.add_paragraph(f"Image not found: {image_path}")
    return doc

def add_section(doc, title, level=1):
    """
    Add a section title to the document.
    """
    return configure_text(doc, title, is_title=True, level=level)

def add_paragraph(doc, text):
    """
    Add a paragraph to the document.
    """
    return configure_text(doc, text)

def analyze_defects(critical_defects):
    """
    Analyze defects and generate insights.
    """
    insights = []

    # Large ERF (> 0.95)
    large_erf_defects = critical_defects[critical_defects['ERF'] > 0.95]
    if not large_erf_defects.empty:
        insights.append(f"{len(large_erf_defects)} defeitos com ERF > 0.95 foram identificados. Estes defeitos são considerados críticos e requerem atenção imediata.")

    # Large defect depth (> 50% of wall thickness)
    large_depth_defects = critical_defects[critical_defects['Depth (d)'] > 0.5 * critical_defects['t']]
    if not large_depth_defects.empty:
        insights.append(f"{len(large_depth_defects)} defeitos com profundidade superior a 50% da espessura da parede foram detectados. Estes defeitos podem comprometer a integridade estrutural do duto.")

    # Defects with high length
    large_length_defects = critical_defects[critical_defects['Length (L)'] > 100]
    if not large_length_defects.empty:
        insights.append(f"{len(large_length_defects)} defeitos com comprimento superior a 100 mm foram identificados. Defeitos longos podem indicar corrosão generalizada.")

    # Defects in clusters
    clustered_defects = critical_defects[critical_defects['Cluster defs'] > 1]
    if not clustered_defects.empty:
        insights.append(f"{len(clustered_defects)} defeitos estão agrupados em clusters. Defeitos interagentes podem aumentar a criticidade e requerem avaliação detalhada.")

    return insights

def Criar_Relatorio(file_name='Relatorio_de_Inspecao.docx', insp=None):
    # Create document and set basic formatting
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # Add title page
    title = doc.add_paragraph()
    title.add_run(f"Relatório de Integridade de Duto\n{insp.date}").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # General Information Section
    doc.add_heading('1. Informações Gerais', level=1)
    general_info = [
        f"Duto: {insp.grid_letter}",
        f"Data da Inspeção: {insp.date}",
        f"Diâmetro Externo (OD): {insp.OD} mm",
        f"MAOP: {insp.MAOP} bar",
        f"Material (σe/σu): {insp.sige} MPa / {insp.sigu} MPa",
        f"Número Total de Defeitos: {len(insp.df_Def)}",
        f"Juntas Inspecionadas: {len(insp.df_joints)}"
    ]
    for info in general_info:
        doc.add_paragraph(info, style='ListBullet')

    # Critical Defects Analysis
    doc.add_heading('2. Análise de Defeitos Críticos', level=1)
    
    # ERF Analysis
    high_risk = insp.df_Def[insp.df_Def['ERF'] > 0.95]
    doc.add_paragraph(
        f"Foram identificados {len(high_risk)} defeitos com ERF > 0.95 (alto risco):",
        style='ListBullet'
    )
    
    # Depth Analysis
    wall_thickness = insp.t0.mean()  # Assuming t0 contains wall thickness values
    depth_threshold = 0.5 * wall_thickness
    deep_defects = insp.df_Def[insp.df_Def['d'] > depth_threshold]
    doc.add_paragraph(
        f"{len(deep_defects)} defeitos com profundidade > 50% da espessura de parede ({depth_threshold:.1f} mm):",
        style='ListBullet'
    )

    # Critical Clusters
    # doc.add_heading('3. Agrupamentos Críticos', level=1)
    # if hasattr(insp, 'crit_cluster') and not insp.crit_cluster.empty:
    #     cluster_info = [
    #         f"Total de agrupamentos críticos: {len(insp.crit_cluster)}",
    #         f"Maior agrupamento: {insp.crit_cluster['L max'].max():.1f} mm",
    #         f"ERF médio em agrupamentos: {insp.crit_cluster['ERF'].mean():.2f}"
    #     ]
    #     for info in cluster_info:
    #         doc.add_paragraph(info, style='ListBullet')
    # else:
    #     doc.add_paragraph("Nenhum agrupamento crítico identificado", style='ListBullet')

    # ERF Distribution Table
    doc.add_heading('4. Distribuição de ERF', level=1)
    table = doc.add_table(insp.ERF_dist.shape[0]+1, insp.ERF_dist.shape[1])
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    for j, column in enumerate(insp.ERF_dist.columns):
        table.cell(0,j).text = str(column)
    
    # Add data rows
    for i in range(insp.ERF_dist.shape[0]):
        for j in range(insp.ERF_dist.shape[1]):
            table.cell(i+1,j).text = str(insp.ERF_dist.iloc[i,j])

    # Future Assessment
    # if hasattr(insp, 'future'):
    #     doc.add_heading('5. Projeção Futura', level=1)
    #     future_text = [
    #         f"Horizonte de análise: {insp.future} anos",
    #         f"Fator de crescimento conservativo: {insp.F}",
    #         "Principais riscos projetados:",
    #         f"- Aumento médio de ERF: {insp.dfg['ERF_increase'].mean():.2f}",
    #         f"- Redução média de MSOP: {insp.dfg['MSOP_reduction'].mean():.2f}%"
    #     ]
    #     for text in future_text:
    #         doc.add_paragraph(text, style='ListBullet')

    # Conclusions and Recommendations
    doc.add_heading('6. Conclusões e Recomendações', level=1)
    conclusions = [
        "Prioridades de manutenção:",
        f"- {len(high_risk)} intervenções imediatas (ERF > 0.95)",
        f"- {len(deep_defects)} reparos prioritários (profundidade crítica)",
        "Recomendações:",
        "- Realizar inspeção detalhada nos clusters identificados",
        "- Monitorar crescimento de defeitos através de ILI periódico",
        "- Revisar programa de proteção catódica"
    ]
    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='ListBullet')

    # Save document
    doc.save(file_name)
    return f"Relatório gerado com sucesso: {file_name}"

def create_report(file_name, inspection_data, critical_defects, figure_paths):
    """
    Generate the pipeline integrity report.
    """
    doc = Document()

    # Report Header
    doc = add_section(doc, "Relatório de Análise de Integridade do Duto & Manual de Manutenção", level=0)
    doc = add_paragraph(doc, f"Duto: {inspection_data.get('pipeline_name', 'Nome do Duto')}")
    doc = add_paragraph(doc, f"Trecho: {inspection_data.get('section', 'Trecho Analisado')}")
    doc = add_paragraph(doc, f"Data da Análise: {date.today()}")
    doc = add_paragraph(doc, f"Inspeção Atual: {inspection_data.get('inspection_date', 'Data da Inspeção')}")
    doc = add_paragraph(doc, f"Comprimento Total Analisado: {inspection_data.get('total_length', 'N/A')} m")
    doc = add_paragraph(doc, f"Número de Juntas Analisadas: {inspection_data.get('joint_count', 'N/A')}")
    doc = add_paragraph(doc, f"Total de Anomalias Detectadas: {inspection_data.get('defect_count', 'N/A')}")

    # Introduction
    doc = add_section(doc, "1. Introdução", level=1)
    doc = add_paragraph(doc, "Este relatório apresenta os resultados da análise de integridade do duto com base nos dados de inspeção.")

    # General Information
    doc = add_section(doc, "2. Informações Gerais", level=1)
    doc = add_paragraph(doc, "Detalhes sobre o duto e a inspeção.")

    # Critical Defects Analysis
    doc = add_section(doc, "3. Análise de Pontos Críticos", level=1)
    insights = analyze_defects(critical_defects)
    if insights:
        for insight in insights:
            doc = add_paragraph(doc, insight)
    else:
        doc = add_paragraph(doc, "Nenhum defeito crítico foi identificado.")

    # Critical Defects Table
    doc = add_section(doc, "4. Tabela de Defeitos Críticos", level=2)
    doc = create_table(doc, critical_defects)

    # Figures and Analysis
    doc = add_section(doc, "5. Análise de Anomalias", level=1)
    for i, figure_path in enumerate(figure_paths):
        doc = add_section(doc, f"Figura {i+1}: {figure_path}", level=2)
        doc = add_image(doc, figure_path)

    # Conclusion
    doc = add_section(doc, "6. Conclusão", level=1)
    doc = add_paragraph(doc, "Resumo das condições do duto e recomendações para manutenção.")

    # Save the document
    doc.save(file_name)

# Example Usage
if __name__ == "__main__":
    inspection_data = {
        "pipeline_name": "Duto XYZ",
        "section": "Trecho A",
        "inspection_date": "2023-10-01",
        "total_length": "1200",
        "joint_count": "150",
        "defect_count": "25"
    }

    critical_defects = pd.DataFrame({
        "Defect ID": [1, 2, 3],
        "Tube Number": [34030.0, 35370.0, 47440.0],
        "Z Position": [34278.05025, 35075.173, 42202.87975],
        "Depth (d)": [30.0, 46.0, 31.0],
        "Length (L)": [129.5, 118.0, 120.5],
        "ERF": [0.936912, 1.055052, 0.93695],
        "Tipo POF": ["AXGR", "GENE", "AXGR"],
        "t": [6.4, 6.4, 6.4],  # Wall thickness
        "Cluster defs": [3.0, 2.0, 2.0]  # Number of defects in cluster
    })

    figure_paths = [
        "anomalies_histogram.png",
        "Depth_histogram.png",
        "Defects_Sizes.png"
    ]

    create_report(
        file_name="Relatorio_de_Inspecao.docx",
        inspection_data=inspection_data,
        critical_defects=critical_defects,
        figure_paths=figure_paths
    )